#!/usr/bin/env python3
"""
ULTIMATE DOCUMENT TRANSLATOR - PRODUCTION VERSION
Features:
- Multiple NMT backends: CTranslate2, NLLB, LLM
- Multiple aligners: Lindat, fast_align, SimAlign, heuristic
- Configurable translation modes with/without alignment
- Production-grade error handling and logging
"""

import argparse
import asyncio
import logging
import os
import re
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Any
from enum import Enum
import requests
from tqdm import tqdm

# Check and import required libraries
print("🔍 Checking libraries...")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.text.paragraph import Paragraph
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
    print("  ✓ python-docx")
except ImportError:
    print("  ✗ python-docx (required)")
    sys.exit(1)

# Optional: CTranslate2
try:
    import ctranslate2
    from transformers import AutoTokenizer
    from huggingface_hub import snapshot_download
    HAS_CT2 = True
    print("  ✓ CTranslate2")
except ImportError:
    HAS_CT2 = False
    print("  ⊘ CTranslate2 (optional)")

# Optional: Transformers for NLLB
try:
    from transformers import AutoModelForSeq2SeqLM, AutoTokenizer as HFTokenizer
    import torch
    HAS_TRANSFORMERS = True
    print("  ✓ Transformers (for NLLB)")
except ImportError:
    HAS_TRANSFORMERS = False
    print("  ⊘ Transformers (optional - needed for NLLB)")

# Optional: OpenAI
try:
    from openai import AsyncOpenAI
    HAS_OPENAI = True
    print("  ✓ OpenAI")
except ImportError:
    HAS_OPENAI = False
    print("  ⊘ OpenAI (optional)")

# Optional: Anthropic
try:
    from anthropic import AsyncAnthropic
    HAS_ANTHROPIC = True
    print("  ✓ Anthropic")
except ImportError:
    HAS_ANTHROPIC = False
    print("  ⊘ Anthropic (optional)")

# Optional: fast_align
try:
    from fast_align import align
    HAS_FAST_ALIGN = True
    print("  ✓ fast_align")
except ImportError:
    HAS_FAST_ALIGN = False
    print("  ⊘ fast_align (optional)")

# Optional: simalign
try:
    from simalign import SentenceAligner
    HAS_SIMALIGN = True
    print("  ✓ simalign")
except ImportError:
    HAS_SIMALIGN = False
    print("  ⊘ simalign (optional)")

print("-" * 60)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============================================================================
# ENUMS FOR CONFIGURATION
# ============================================================================

class TranslationBackend(Enum):
    """Available translation backends"""
    CT2 = "ct2"
    NLLB = "nllb"
    OPENAI = "openai"
    ANTHROPIC = "anthropic"
    OLLAMA = "ollama"
    AUTO = "auto"


class AlignerBackend(Enum):
    """Available alignment backends"""
    LINDAT = "lindat"
    FAST_ALIGN = "fast_align"
    SIMALIGN = "simalign"
    HEURISTIC = "heuristic"
    AUTO = "auto"


class TranslationMode(Enum):
    """Translation modes"""
    NMT_ONLY = "nmt"  # Fast NMT without alignment
    LLM_WITH_ALIGN = "llm-align"  # LLM + alignment for formatting
    LLM_WITHOUT_ALIGN = "llm-plain"  # LLM without alignment
    HYBRID = "hybrid"  # NMT with LLM fallback + alignment


# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class FormatRun:
    """A run of text with formatting"""
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    font_color: Optional[Tuple[int, int, int]] = None


@dataclass
class TranslatableParagraph:
    """Paragraph with formatting and metadata"""
    runs: List[FormatRun] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def get_text(self) -> str:
        return ''.join(run.text for run in self.runs)
    
    def get_words(self) -> List[str]:
        """Get words for alignment"""
        text = self.get_text()
        return text.split()
    
    def get_formatted_word_indices(self) -> Dict[str, Set[int]]:
        """Extract which word indices have which formatting."""
        formatted = {'italic': set(), 'bold': set(), 'italic_bold': set()}
        
        text = self.get_text()
        words = text.split()
        
        if not words:
            return formatted
        
        # Build character-to-word map
        char_to_word = {}
        char_pos = 0
        for word_idx, word in enumerate(words):
            while char_pos < len(text) and text[char_pos].isspace():
                char_pos += 1
            for i in range(len(word)):
                if char_pos < len(text):
                    char_to_word[char_pos] = word_idx
                    char_pos += 1
        
        # Check which words are formatted
        char_pos = 0
        for run in self.runs:
            if not run.text:
                continue
            
            for char in run.text:
                if char_pos in char_to_word:
                    word_idx = char_to_word[char_pos]
                    
                    if not char.isspace():
                        if run.bold and run.italic:
                            formatted['italic_bold'].add(word_idx)
                        elif run.italic:
                            formatted['italic'].add(word_idx)
                        elif run.bold:
                            formatted['bold'].add(word_idx)
                
                char_pos += 1
        
        return formatted


# ============================================================================
# TRANSLATION BACKENDS
# ============================================================================

class CTranslate2Translator:
    """CTranslate2 for fast translation"""
    
    MODELS = {
        'en_to_x': 'cstr/wmt21ct2_int8',
        'x_to_en': 'cstr/wmt21-ct2-x-en-int8',
    }
    
    SUPPORTED_LANGS = ['de', 'es', 'fr', 'it', 'ja', 'zh', 'ru', 'pt', 'nl', 'cs', 'uk']
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = False
        
        if not HAS_CT2:
            logger.info("⊘ CT2: library not installed")
            return
        
        if src_lang == 'en' and tgt_lang in self.SUPPORTED_LANGS:
            self.direction = 'en_to_x'
            self.tokenizer_name = "facebook/wmt21-dense-24-wide-en-x"
        elif tgt_lang == 'en' and src_lang in self.SUPPORTED_LANGS:
            self.direction = 'x_to_en'
            self.tokenizer_name = "facebook/wmt21-dense-24-wide-x-en"
        else:
            logger.info(f"⊘ CT2: {src_lang}→{tgt_lang} not supported")
            return
        
        try:
            logger.info(f"Loading CT2 model...")
            model_path = self._get_or_download_model()
            self.translator = ctranslate2.Translator(model_path, device="auto")
            self.tokenizer = AutoTokenizer.from_pretrained(self.tokenizer_name)
            self.tokenizer.src_lang = src_lang
            self.tokenizer.tgt_lang = tgt_lang
            self.available = True
            logger.info(f"✓ CT2 ready ({src_lang}→{tgt_lang})")
        except Exception as e:
            logger.error(f"CT2 init failed: {e}")
    
    def _get_or_download_model(self) -> str:
        cache_base = Path.home() / '.cache' / 'huggingface' / 'hub'
        model_repo = self.MODELS[self.direction]
        model_dir = cache_base / f"models--{model_repo.replace('/', '--')}"
        ref_path = model_dir / 'refs' / 'main'
        
        if ref_path.exists():
            with open(ref_path) as f:
                commit_hash = f.read().strip()
            snapshot_path = model_dir / 'snapshots' / commit_hash
            if (snapshot_path / 'model.bin').exists():
                return str(snapshot_path)
        
        logger.info("Downloading CT2 model...")
        return snapshot_download(repo_id=model_repo)
    
    def translate_batch(self, texts: List[str]) -> List[str]:
        """Translate batch of texts"""
        if not self.available or not texts:
            return texts
        
        try:
            source_batches = []
            for text in texts:
                if not text.strip():
                    source_batches.append([self.tokenizer.src_lang])
                    continue
                
                tokens = self.tokenizer.tokenize(text)
                if self.tokenizer.src_lang not in tokens:
                    tokens = [self.tokenizer.src_lang] + tokens + [self.tokenizer.eos_token]
                source_batches.append(tokens)
            
            target_prefix = [self.tokenizer.lang_code_to_token[self.tgt_lang]]
            results = self.translator.translate_batch(
                source_batches,
                target_prefix=[target_prefix] * len(texts),
                beam_size=5,
                repetition_penalty=1.5
            )
            
            translated = []
            for i, res in enumerate(results):
                if not texts[i].strip():
                    translated.append(texts[i])
                    continue
                
                target_tokens = res.hypotheses[0][len(target_prefix):]
                decoded = self.tokenizer.decode(
                    self.tokenizer.convert_tokens_to_ids(target_tokens),
                    skip_special_tokens=True
                )
                translated.append(decoded.strip())
            
            return translated
        except Exception as e:
            logger.error(f"CT2 translation failed: {e}")
            return texts


class NLLBTranslator:
    """NLLB-200 translator for 200+ languages"""
    
    # Language code mapping for NLLB
    LANG_CODES = {
        'en': 'eng_Latn', 'de': 'deu_Latn', 'fr': 'fra_Latn', 
        'es': 'spa_Latn', 'it': 'ita_Latn', 'pt': 'por_Latn',
        'ru': 'rus_Cyrl', 'zh': 'zho_Hans', 'ja': 'jpn_Jpan',
        'ko': 'kor_Hang', 'ar': 'arb_Arab', 'hi': 'hin_Deva',
        'nl': 'nld_Latn', 'pl': 'pol_Latn', 'tr': 'tur_Latn',
        'cs': 'ces_Latn', 'uk': 'ukr_Cyrl', 'vi': 'vie_Latn',
    }
    
    def __init__(self, src_lang: str, tgt_lang: str, model_size: str = "600M"):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = False
        
        if not HAS_TRANSFORMERS:
            logger.info("⊘ NLLB: transformers not installed")
            return
        
        # Map to NLLB codes
        self.src_code = self.LANG_CODES.get(src_lang)
        self.tgt_code = self.LANG_CODES.get(tgt_lang)
        
        if not self.src_code or not self.tgt_code:
            logger.info(f"⊘ NLLB: {src_lang}→{tgt_lang} not in language map")
            return
        
        try:
            logger.info(f"Loading NLLB-{model_size} model...")
            model_name = f"facebook/nllb-200-{model_size}"
            
            self.tokenizer = HFTokenizer.from_pretrained(model_name)
            self.model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
            
            # Move to GPU if available
            self.device = "cuda" if torch.cuda.is_available() else "cpu"
            self.model.to(self.device)
            
            self.available = True
            logger.info(f"✓ NLLB-{model_size} ready on {self.device} ({src_lang}→{tgt_lang})")
        except Exception as e:
            logger.error(f"NLLB init failed: {e}")
    
    def translate_batch(self, texts: List[str], batch_size: int = 8) -> List[str]:
        """Translate batch of texts"""
        if not self.available or not texts:
            return texts
        
        try:
            results = []
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                
                # Tokenize with forced source language
                inputs = self.tokenizer(
                    batch, 
                    return_tensors="pt", 
                    padding=True, 
                    truncation=True,
                    max_length=512
                ).to(self.device)
                
                # Generate with forced target language
                translated_tokens = self.model.generate(
                    **inputs,
                    forced_bos_token_id=self.tokenizer.lang_code_to_id[self.tgt_code],
                    max_length=512,
                    num_beams=5,
                    early_stopping=True
                )
                
                # Decode
                batch_results = self.tokenizer.batch_decode(
                    translated_tokens, 
                    skip_special_tokens=True
                )
                results.extend(batch_results)
            
            return results
        except Exception as e:
            logger.error(f"NLLB translation failed: {e}")
            return texts


class LLMTranslator:
    """LLM translator (OpenAI/Anthropic/Ollama)"""
    
    def __init__(self, src_lang: str, tgt_lang: str, preferred_provider: Optional[str] = None):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.providers = self._init_providers(preferred_provider)
        
        if self.providers:
            logger.info(f"✓ LLM available ({list(self.providers.keys())})")
    
    def _init_providers(self, preferred: Optional[str] = None) -> Dict[str, Any]:
        providers = {}
        
        # OpenAI
        if HAS_OPENAI and os.getenv("OPENAI_API_KEY"):
            providers["openai"] = {
                "client": AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY")),
                "model": os.getenv("OPENAI_MODEL", "gpt-4o-mini")
            }
        
        # Anthropic
        if HAS_ANTHROPIC and os.getenv("ANTHROPIC_API_KEY"):
            providers["anthropic"] = {
                "client": AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY")),
                "model": os.getenv("ANTHROPIC_MODEL", "claude-3-5-sonnet-20241022")
            }
        
        # Ollama
        if self._check_ollama():
            providers["ollama"] = {
                "url": "http://localhost:11434/api/generate",
                "model": self._get_ollama_model()
            }
        
        # Filter to preferred if specified
        if preferred and preferred in providers:
            return {preferred: providers[preferred]}
        
        return providers
    
    def _check_ollama(self) -> bool:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            return r.status_code == 200 and len(r.json().get("models", [])) > 0
        except:
            return False
    
    def _get_ollama_model(self) -> str:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            models = r.json().get("models", [])
            return models[0]["name"] if models else "llama3.2"
        except:
            return "llama3.2"
    
    async def translate_text(self, text: str, use_alignment: bool = True) -> Optional[str]:
        """Translate single text"""
        if not text.strip() or not self.providers:
            return None
        
        # Adjust prompt based on alignment mode
        if use_alignment:
            prompt = (
                f"Translate the following text from {self.src_lang} to {self.tgt_lang}. "
                f"Preserve the word order as much as possible for alignment purposes. "
                f"Return ONLY the translation:\n\n{text}"
            )
        else:
            prompt = (
                f"Translate the following text from {self.src_lang} to {self.tgt_lang}. "
                f"Provide a natural, fluent translation. "
                f"Return ONLY the translation:\n\n{text}"
            )
        
        # Try providers in order
        for provider_name, provider in self.providers.items():
            try:
                if provider_name == "openai":
                    response = await provider["client"].chat.completions.create(
                        model=provider["model"],
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                        max_tokens=4000
                    )
                    return response.choices[0].message.content.strip()
                
                elif provider_name == "anthropic":
                    response = await provider["client"].messages.create(
                        model=provider["model"],
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                        max_tokens=4000
                    )
                    return response.content[0].text.strip()
                
                elif provider_name == "ollama":
                    r = requests.post(
                        provider["url"],
                        json={"model": provider["model"], "prompt": prompt, "stream": False},
                        timeout=120
                    )
                    if r.status_code == 200:
                        return r.json().get("response", "").strip()
            
            except Exception as e:
                logger.debug(f"{provider_name} failed: {e}")
                continue
        
        return None
    
    async def translate_batch(self, texts: List[str], use_alignment: bool = True) -> List[str]:
        """Translate batch"""
        tasks = [self.translate_text(text, use_alignment) for text in texts]
        results = await asyncio.gather(*tasks)
        return [res if res else text for res, text in zip(results, texts)]


# ============================================================================
# WORD ALIGNERS
# ============================================================================

class LindatAligner:
    """Lindat word alignment API"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = self._check_available()
        if self.available:
            logger.info(f"✓ Lindat aligner available ({src_lang}-{tgt_lang})")
    
    def _check_available(self) -> bool:
        try:
            r = requests.get(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                timeout=5
            )
            return r.status_code in [200, 405]
        except:
            return False
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            r = requests.post(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                headers={'Content-Type': 'application/json'},
                json={'src_tokens': [src_words], 'trg_tokens': [tgt_words]},
                timeout=30
            )
            
            if r.status_code == 200:
                alignment = r.json()["alignment"][0]
                return [(int(a[0]), int(a[1])) for a in alignment]
        except Exception as e:
            logger.debug(f"Lindat alignment failed: {e}")
        
        return []


class FastAlignAligner:
    """fast_align local aligner"""
    
    def __init__(self):
        self.available = HAS_FAST_ALIGN
        if self.available:
            logger.info("✓ fast_align available")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            src_text = ' '.join(src_words)
            tgt_text = ' '.join(tgt_words)
            alignment_text = f"{src_text} ||| {tgt_text}"
            
            result = align([alignment_text], forward=True)
            
            alignments = []
            for align_str in result[0].split():
                src_idx, tgt_idx = map(int, align_str.split('-'))
                alignments.append((src_idx, tgt_idx))
            
            return alignments
        except Exception as e:
            logger.debug(f"fast_align failed: {e}")
            return []


class SimAlignAligner:
    """SimAlign aligner"""
    
    def __init__(self):
        self.available = HAS_SIMALIGN
        if self.available:
            try:
                self.aligner = SentenceAligner(model="bert", token_type="bpe", matching_methods="mai")
                logger.info("✓ SimAlign available")
            except Exception as e:
                logger.debug(f"SimAlign init failed: {e}")
                self.available = False
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            result = self.aligner.get_word_aligns(src_words, tgt_words)
            alignments = []
            for src_idx, tgt_idx in result['mwmf'].items():
                alignments.append((src_idx, tgt_idx))
            return alignments
        except Exception as e:
            logger.debug(f"SimAlign failed: {e}")
            return []


class HeuristicAligner:
    """Heuristic fallback - align words that appear in both"""
    
    def __init__(self):
        logger.info("✓ Heuristic aligner (fallback)")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Simple heuristic alignment"""
        alignments = []
        
        src_lower = [w.lower().strip('.,!?;:') for w in src_words]
        tgt_lower = [w.lower().strip('.,!?;:') for w in tgt_words]
        
        for i, src_word in enumerate(src_lower):
            for j, tgt_word in enumerate(tgt_lower):
                if src_word == tgt_word and len(src_word) > 2:
                    alignments.append((i, j))
        
        return alignments


class MultiAligner:
    """Try multiple aligners in sequence"""
    
    def __init__(self, src_lang: str, tgt_lang: str, preferred: Optional[str] = None):
        self.aligners = []
        
        # Build aligner list based on preference
        if preferred == "lindat" or preferred == "auto":
            lindat = LindatAligner(src_lang, tgt_lang)
            if lindat.available:
                self.aligners.append(("Lindat", lindat))
        
        if preferred == "fast_align" or preferred == "auto":
            fast_align = FastAlignAligner()
            if fast_align.available:
                self.aligners.append(("fast_align", fast_align))
        
        if preferred == "simalign" or preferred == "auto":
            simalign = SimAlignAligner()
            if simalign.available:
                self.aligners.append(("SimAlign", simalign))
        
        # Always include heuristic as fallback
        if preferred == "heuristic" or preferred == "auto":
            self.aligners.append(("Heuristic", HeuristicAligner()))
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Try aligners in order until one succeeds"""
        for name, aligner in self.aligners:
            result = aligner.align(src_words, tgt_words)
            if result:
                logger.debug(f"Using {name} alignment: {len(result)} links")
                return result
        
        logger.warning("All aligners failed, returning empty alignment")
        return []


# ============================================================================
# DOCUMENT TRANSLATOR
# ============================================================================

class UltimateDocumentTranslator:
    """Production-grade document translator with configurable backends"""
    
    def __init__(
        self,
        src_lang: str,
        tgt_lang: str,
        mode: TranslationMode = TranslationMode.HYBRID,
        nmt_backend: Optional[str] = None,
        llm_provider: Optional[str] = None,
        aligner: Optional[str] = None,
        nllb_model_size: str = "600M"
    ):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.mode = mode
        
        logger.info(f"Initializing translator ({src_lang}→{tgt_lang}, mode={mode.value})")
        
        # Initialize translators based on mode
        self.ct2 = None
        self.nllb = None
        self.llm = None
        
        if mode in [TranslationMode.NMT_ONLY, TranslationMode.HYBRID]:
            # Try NMT backends
            if nmt_backend == "nllb" or nmt_backend == "auto":
                self.nllb = NLLBTranslator(src_lang, tgt_lang, nllb_model_size)
            
            if nmt_backend == "ct2" or nmt_backend == "auto":
                self.ct2 = CTranslate2Translator(src_lang, tgt_lang)
        
        if mode in [TranslationMode.LLM_WITH_ALIGN, TranslationMode.LLM_WITHOUT_ALIGN, TranslationMode.HYBRID]:
            self.llm = LLMTranslator(src_lang, tgt_lang, llm_provider)
        
        # Initialize aligner if needed
        self.aligner = None
        if mode in [TranslationMode.LLM_WITH_ALIGN, TranslationMode.HYBRID]:
            self.aligner = MultiAligner(src_lang, tgt_lang, aligner or "auto")
        
        # Check we have at least one backend
        has_backend = False
        if self.ct2 and self.ct2.available:
            has_backend = True
            logger.info("Primary backend: CTranslate2")
        if self.nllb and self.nllb.available:
            has_backend = True
            logger.info("Primary backend: NLLB")
        if self.llm and self.llm.providers:
            has_backend = True
            logger.info(f"LLM backend: {list(self.llm.providers.keys())}")
        
        if not has_backend:
            logger.error("No translation backends available!")
            sys.exit(1)
    
    async def translate_text(self, text: str) -> str:
        """Translate single text with fallback chain"""
        if not text.strip():
            return text
        
        # Choose translation based on mode
        if self.mode == TranslationMode.NMT_ONLY:
            # Try NMT backends only
            if self.nllb and self.nllb.available:
                result = self.nllb.translate_batch([text])[0]
                if result and result != text:
                    return result
            
            if self.ct2 and self.ct2.available:
                result = self.ct2.translate_batch([text])[0]
                if result and result != text:
                    return result
        
        elif self.mode == TranslationMode.LLM_WITH_ALIGN:
            # Use LLM with alignment-friendly prompts
            if self.llm and self.llm.providers:
                result = await self.llm.translate_text(text, use_alignment=True)
                if result:
                    return result
        
        elif self.mode == TranslationMode.LLM_WITHOUT_ALIGN:
            # Use LLM with natural translation
            if self.llm and self.llm.providers:
                result = await self.llm.translate_text(text, use_alignment=False)
                if result:
                    return result
        
        elif self.mode == TranslationMode.HYBRID:
            # Try NMT first, then LLM as fallback
            if self.nllb and self.nllb.available:
                result = self.nllb.translate_batch([text])[0]
                if result and result != text:
                    return result
            
            if self.ct2 and self.ct2.available:
                result = self.ct2.translate_batch([text])[0]
                if result and result != text:
                    return result
            
            if self.llm and self.llm.providers:
                result = await self.llm.translate_text(text, use_alignment=True)
                if result:
                    return result
        
        logger.error(f"Translation failed for: {text[:50]}...")
        return text
    
    def extract_paragraph(self, para: Paragraph) -> TranslatableParagraph:
        """Extract paragraph with style and formatting metadata."""
        runs = []
        for run in para.runs:
            font_color = None
            if run.font.color and run.font.color.rgb:
                try:
                    rgb = run.font.color.rgb
                    font_color = (rgb[0], rgb[1], rgb[2])
                except: 
                    pass
            
            runs.append(FormatRun(
                text=run.text,
                bold=run.font.bold,
                italic=run.font.italic,
                underline=run.font.underline,
                font_name=run.font.name,
                font_size=run.font.size.pt if run.font.size else None,
                font_color=font_color
            ))
        
        trans_para = TranslatableParagraph(runs=runs)
        trans_para.metadata['style'] = para.style
        trans_para.metadata['alignment'] = para.alignment
        
        pf = para.paragraph_format
        trans_para.metadata['layout'] = {
            'left_indent': pf.left_indent,
            'right_indent': pf.right_indent,
            'first_line_indent': pf.first_line_indent,
            'line_spacing': pf.line_spacing,
            'space_before': pf.space_before,
            'space_after': pf.space_after
        }
        return trans_para
    
    def copy_font_properties(self, target_run, source_run: FormatRun):
        """Copies all font properties including safety checks."""
        if source_run.font_name:
            target_run.font.name = source_run.font_name
        if source_run.font_size:
            target_run.font.size = Pt(source_run.font_size)
        if source_run.font_color:
            target_run.font.color.rgb = RGBColor(*source_run.font_color)
        if source_run.underline is not None:
            target_run.font.underline = source_run.underline
    
    def apply_aligned_formatting(
        self, 
        para: Paragraph, 
        trans_para: TranslatableParagraph,
        translated_text: str, 
        alignment: List[Tuple[int, int]]
    ):
        """Applies formatting while preserving paragraph layout and footnotes."""
        formatted_words = trans_para.get_formatted_word_indices()
        src_words = trans_para.get_words()
        tgt_words = translated_text.split()
        
        if not tgt_words:
            return
        
        # Target formatting map
        tgt_formatting = {}
        for src_idx, tgt_idx in alignment:
            if src_idx < len(src_words) and tgt_idx < len(tgt_words):
                for f_type in ['italic_bold', 'bold', 'italic']:
                    if src_idx in formatted_words[f_type]:
                        tgt_formatting[tgt_idx] = f_type
                        break
        
        # Preserve footnote references
        p = para._p
        for run_element in p.xpath('.//w:r'):
            if not run_element.xpath('.//w:footnoteRef | .//w:footnoteReference'):
                run_element.getparent().remove(run_element)
        
        # Restore paragraph layout
        para.alignment = trans_para.metadata['alignment']
        if trans_para.metadata['style']:
            para.style = trans_para.metadata['style']
        
        layout = trans_para.metadata['layout']
        para.paragraph_format.left_indent = layout['left_indent']
        para.paragraph_format.first_line_indent = layout['first_line_indent']
        para.paragraph_format.line_spacing = layout['line_spacing']
        
        # Create runs with formatting
        font_template = next((r for r in trans_para.runs if r.text.strip()), None)
        
        runs_to_create = []
        current_format = tgt_formatting.get(0, None)
        current_words = [tgt_words[0]]
        
        for i in range(1, len(tgt_words)):
            fmt = tgt_formatting.get(i, None)
            if fmt == current_format:
                current_words.append(tgt_words[i])
            else:
                runs_to_create.append((current_format, current_words))
                current_format, current_words = fmt, [tgt_words[i]]
        runs_to_create.append((current_format, current_words))
        
        for i, (f_type, words) in enumerate(runs_to_create):
            text = ' '.join(words) + (" " if i < len(runs_to_create) - 1 else "")
            run = para.add_run(text)
            
            if f_type == 'italic_bold':
                run.bold = run.italic = True
            elif f_type == 'bold':
                run.bold = True
            elif f_type == 'italic':
                run.italic = True
            
            if font_template:
                self.copy_font_properties(run, font_template)
    
    def apply_plain_formatting(self, para: Paragraph, trans_para: TranslatableParagraph, translated_text: str):
        """Apply formatting without alignment (for LLM_WITHOUT_ALIGN mode)"""
        # Clear existing runs except footnotes
        p = para._p
        for run_element in p.xpath('.//w:r'):
            if not run_element.xpath('.//w:footnoteRef | .//w:footnoteReference'):
                run_element.getparent().remove(run_element)
        
        # Restore paragraph layout
        para.alignment = trans_para.metadata['alignment']
        if trans_para.metadata['style']:
            para.style = trans_para.metadata['style']
        
        layout = trans_para.metadata['layout']
        para.paragraph_format.left_indent = layout['left_indent']
        para.paragraph_format.first_line_indent = layout['first_line_indent']
        para.paragraph_format.line_spacing = layout['line_spacing']
        
        # Add translated text with default formatting from first run
        font_template = next((r for r in trans_para.runs if r.text.strip()), None)
        run = para.add_run(translated_text)
        
        if font_template:
            self.copy_font_properties(run, font_template)
    
    def is_paragraph_safe_to_translate(self, para: Paragraph) -> bool:
        """Check if paragraph can be safely translated"""
        if not para.text or not para.text.strip():
            return False
        
        if len(para.text.strip()) <= 1 and not any(run.text.strip() for run in para.runs):
            return False
        
        try:
            for run in para.runs:
                if run._element.xpath('.//w:drawing | .//w:pict'):
                    logger.debug(f"Skipping paragraph with drawing/image")
                    return False
        except:
            pass
        
        try:
            field_chars = para._element.xpath('.//w:fldChar')
            if field_chars:
                is_footnote_field = para._element.xpath('.//w:footnoteReference')
                if not is_footnote_field:
                    logger.debug(f"Skipping paragraph with field")
                    return False
        except:
            pass
        
        return True
    
    async def translate_paragraph(self, para: Paragraph):
        """Translate single paragraph"""
        if not para.text or not para.text.strip():
            return
        
        if not self.is_paragraph_safe_to_translate(para):
            return
        
        try:
            trans_para = self.extract_paragraph(para)
            original_text = trans_para.get_text()
            
            if not original_text.strip():
                return
            
            # Translate
            translated_text = await self.translate_text(original_text)
            
            if not translated_text.strip():
                return
            
            # Apply formatting based on mode
            if self.mode == TranslationMode.LLM_WITHOUT_ALIGN:
                # No alignment, just apply plain formatting
                self.apply_plain_formatting(para, trans_para, translated_text)
            else:
                # Use alignment
                src_words = trans_para.get_words()
                tgt_words = translated_text.split()
                
                if not src_words or not tgt_words:
                    return
                
                alignment = []
                if self.aligner:
                    alignment = self.aligner.align(src_words, tgt_words)
                
                logger.debug(f"SRC: {src_words}")
                logger.debug(f"TGT: {tgt_words}")
                logger.debug(f"ALIGN: {alignment}")
                
                self.apply_aligned_formatting(para, trans_para, translated_text, alignment)
            
        except Exception as e:
            logger.error(f"Failed to translate paragraph: {e}")
            logger.debug(f"  Text: {para.text[:100]}...")
    
    def get_footnotes(self, doc: Document) -> List[Paragraph]:
        """Extract footnotes from document"""
        document_part = doc.part
        footnote_part = None
        
        for rel in document_part.rels.values():
            if "relationships/footnotes" in rel.reltype:
                footnote_part = rel.target_part
                break
        
        if not footnote_part:
            try:
                import zipfile
                with zipfile.ZipFile(doc._part._package.blob) as z:
                    if 'word/footnotes.xml' in z.namelist():
                        logger.info("✓ Found footnotes via ZIP")
            except:
                pass
            return []
        
        from docx.oxml import parse_xml
        root = parse_xml(footnote_part.blob)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        paragraphs = []
        for footnote in root.xpath('//w:footnote', namespaces=ns):
            f_id = footnote.get(f'{{{ns["w"]}}}id')
            if f_id and int(f_id) <= 0:
                continue
            
            for p_elem in footnote.xpath('.//w:p', namespaces=ns):
                para = Paragraph(p_elem, footnote_part)
                if para.text.strip():
                    paragraphs.append(para)
        
        self._footnote_root = root
        self._footnote_part = footnote_part
        
        return paragraphs
    
    def get_all_paragraphs(self, doc: Document) -> List[Tuple[Paragraph, str]]:
        """Aggregate all paragraphs from document"""
        all_paras = []
        
        # Main body
        for para in doc.paragraphs:
            all_paras.append((para, "body"))
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_paras.append((para, "table"))
        
        # Footnotes
        footnote_paras = self.get_footnotes(doc)
        for para in footnote_paras:
            all_paras.append((para, "footnote"))
        
        # Headers/Footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                all_paras.append((para, "header"))
            for para in section.footer.paragraphs:
                all_paras.append((para, "footer"))
        
        return all_paras
    
    async def translate_document(self, input_path: Path, output_path: Path):
        """Translate entire document"""
        logger.info(f"Loading document: {input_path}")
        
        try:
            doc = Document(str(input_path))
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            return
        
        all_paras_with_location = self.get_all_paragraphs(doc)
        
        translatable = []
        for para, location in all_paras_with_location:
            if para.text and para.text.strip() and self.is_paragraph_safe_to_translate(para):
                translatable.append((para, location))
        
        by_location = defaultdict(int)
        for _, location in translatable:
            by_location[location] += 1
        
        logger.info(f"Found {len(translatable)} paragraphs to translate:")
        for location, count in by_location.items():
            logger.info(f"  - {location}: {count}")
        
        for para, location in tqdm(translatable, desc="Translating"):
            try:
                await self.translate_paragraph(para)
                await asyncio.sleep(0.05)
            except Exception as e:
                logger.error(f"Error translating {location} paragraph: {e}")
                continue
        
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            if output_path.exists():
                output_path.unlink()
            
            logger.info(f"Saving to: {output_path}")
            
            if hasattr(self, '_footnote_part') and hasattr(self, '_footnote_root'):
                from lxml import etree
                self._footnote_part._blob = etree.tostring(self._footnote_root)
            
            doc.save(str(output_path))
            
            if not output_path.exists():
                raise Exception("Output file was not created")
            
            if output_path.stat().st_size == 0:
                raise Exception("Output file is empty")
            
            test_doc = Document(str(output_path))
            logger.info(f"✓ Document verified ({len(test_doc.paragraphs)} paragraphs)")
            logger.info("✓ Translation complete!")
            
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            if output_path.exists():
                output_path.unlink()
            raise


# ============================================================================
# CLI
# ============================================================================

async def main():
    parser = argparse.ArgumentParser(
        description='Document Translator',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Default (auto-detect best backends)
  %(prog)s input.docx output.docx -s en -t de
  
  # Use NLLB for translation
  %(prog)s input.docx output.docx -s en -t fr --nmt nllb
  
  # Use LLM with alignment
  %(prog)s input.docx output.docx -s en -t es --mode llm-align --llm openai
  
  # Use LLM without alignment (natural translation)
  %(prog)s input.docx output.docx -s en -t it --mode llm-plain --llm anthropic
  
  # Hybrid mode with specific aligner
  %(prog)s input.docx output.docx -s en -t de --mode hybrid --aligner fast_align
  
  # Use larger NLLB model
  %(prog)s input.docx output.docx -s en -t ja --nmt nllb --nllb-size 1.3B

Environment Variables:
  OPENAI_API_KEY       - OpenAI API key
  OPENAI_MODEL         - OpenAI model (default: gpt-4o-mini)
  ANTHROPIC_API_KEY    - Anthropic API key
  ANTHROPIC_MODEL      - Anthropic model (default: claude-3-5-sonnet-20241022)
        """
    )
    
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('output', help='Output .docx file')
    parser.add_argument('-s', '--source', default='en', help='Source language (default: en)')
    parser.add_argument('-t', '--target', default='de', help='Target language (default: de)')
    
    parser.add_argument(
        '--mode',
        choices=['nmt', 'llm-align', 'llm-plain', 'hybrid'],
        default='hybrid',
        help='Translation mode (default: hybrid)'
    )
    
    parser.add_argument(
        '--nmt',
        choices=['ct2', 'nllb', 'auto'],
        help='NMT backend (default: auto)'
    )
    
    parser.add_argument(
        '--nllb-size',
        choices=['600M', '1.3B', '3.3B'],
        default='600M',
        help='NLLB model size (default: 600M)'
    )
    
    parser.add_argument(
        '--llm',
        choices=['openai', 'anthropic', 'ollama'],
        help='LLM provider (default: auto-detect)'
    )
    
    parser.add_argument(
        '--aligner',
        choices=['lindat', 'fast_align', 'simalign', 'heuristic', 'auto'],
        help='Word aligner (default: auto)'
    )
    
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"Error: Only .docx files supported")
        sys.exit(1)
    
    mode_map = {
        'nmt': TranslationMode.NMT_ONLY,
        'llm-align': TranslationMode.LLM_WITH_ALIGN,
        'llm-plain': TranslationMode.LLM_WITHOUT_ALIGN,
        'hybrid': TranslationMode.HYBRID
    }
    
    print(f"\n{'='*60}")
    print(f"Ultimate Document Translator - Production Version")
    print(f"{'='*60}")
    print(f"Input:      {input_path}")
    print(f"Output:     {args.output}")
    print(f"Direction:  {args.source.upper()} → {args.target.upper()}")
    print(f"Mode:       {args.mode}")
    if args.nmt:
        print(f"NMT:        {args.nmt}")
    if args.nmt == 'nllb':
        print(f"NLLB Size:  {args.nllb_size}")
    if args.llm:
        print(f"LLM:        {args.llm}")
    if args.aligner:
        print(f"Aligner:    {args.aligner}")
    print(f"{'='*60}\n")
    
    translator = UltimateDocumentTranslator(
        src_lang=args.source,
        tgt_lang=args.target,
        mode=mode_map[args.mode],
        nmt_backend=args.nmt,
        llm_provider=args.llm,
        aligner=args.aligner,
        nllb_model_size=args.nllb_size
    )
    
    await translator.translate_document(input_path, Path(args.output))
    
    print(f"\n{'='*60}")
    print(f"✓ Success!")
    print(f"Output saved to: {args.output}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    asyncio.run(main())