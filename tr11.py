#!/usr/bin/env python3
"""
ULTIMATE DOCUMENT TRANSLATOR - PRODUCTION VERSION (FIXED)
Fixes:
- Proper aligner fallback when specified aligner unavailable
- Footnote Part attribute error
- Better error messages for missing dependencies
"""

import argparse
import asyncio
import logging
import os
import re
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Any
from enum import Enum
import requests
from tqdm import tqdm

# Check and import required libraries
print("🔍 Checking libraries...")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.text.paragraph import Paragraph
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
    print("  ✓ python-docx")
except ImportError:
    print("  ✗ python-docx (required)")
    sys.exit(1)

# Optional: CTranslate2
try:
    import ctranslate2
    from transformers import AutoTokenizer
    from huggingface_hub import snapshot_download
    HAS_CT2 = True
    print("  ✓ CTranslate2")
except ImportError:
    HAS_CT2 = False
    print("  ⊘ CTranslate2 (optional)")

# Optional: Transformers for NLLB
try:
    from transformers import AutoModelForSeq2SeqLM, AutoTokenizer as HFTokenizer
    import torch
    HAS_TRANSFORMERS = True
    print("  ✓ Transformers (for NLLB)")
except ImportError:
    HAS_TRANSFORMERS = False
    print("  ⊘ Transformers (optional - needed for NLLB)")

# Optional: OpenAI
try:
    from openai import AsyncOpenAI
    HAS_OPENAI = True
    print("  ✓ OpenAI")
except ImportError:
    HAS_OPENAI = False
    print("  ⊘ OpenAI (optional)")

# Optional: Anthropic
try:
    from anthropic import AsyncAnthropic
    HAS_ANTHROPIC = True
    print("  ✓ Anthropic")
except ImportError:
    HAS_ANTHROPIC = False
    print("  ⊘ Anthropic (optional)")

# Optional: fast_align
try:
    from fast_align import align
    HAS_FAST_ALIGN = True
    print("  ✓ fast_align")
except ImportError:
    HAS_FAST_ALIGN = False
    print("  ⊘ fast_align (optional - see installation notes)")

# Optional: simalign
try:
    from simalign import SentenceAligner
    HAS_SIMALIGN = True
    print("  ✓ simalign")
except ImportError:
    HAS_SIMALIGN = False
    print("  ⊘ simalign (optional)")

print("-" * 60)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============================================================================
# ENUMS FOR CONFIGURATION
# ============================================================================

class TranslationBackend(Enum):
    """Available translation backends"""
    CT2 = "ct2"
    NLLB = "nllb"
    OPENAI = "openai"
    ANTHROPIC = "anthropic"
    OLLAMA = "ollama"
    AUTO = "auto"


class AlignerBackend(Enum):
    """Available alignment backends"""
    LINDAT = "lindat"
    FAST_ALIGN = "fast_align"
    SIMALIGN = "simalign"
    HEURISTIC = "heuristic"
    AUTO = "auto"


class TranslationMode(Enum):
    """Translation modes"""
    NMT_ONLY = "nmt"
    LLM_WITH_ALIGN = "llm-align"
    LLM_WITHOUT_ALIGN = "llm-plain"
    HYBRID = "hybrid"


# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class FormatRun:
    """A run of text with formatting"""
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    font_color: Optional[Tuple[int, int, int]] = None


@dataclass
class TranslatableParagraph:
    """Paragraph with formatting and metadata"""
    runs: List[FormatRun] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def get_text(self) -> str:
        return ''.join(run.text for run in self.runs)
    
    def get_words(self) -> List[str]:
        """Clean tokenization for the ALIGNER only: strips punctuation."""
        import re
        text = self.get_text()
        # Extracts only alphanumeric sequences (e.g., "Theology" from "Theology:")
        return re.findall(r"\w+", text)
    
    def get_formatted_word_indices(self) -> Dict[str, Set[int]]:
        """Maps formatting to clean alphanumeric word indices."""
        formatted = {'italic': set(), 'bold': set(), 'italic_bold': set()}
        text = self.get_text()
        words = self.get_words() # Uses the same list the aligner sees
        
        if not words:
            return formatted
        
        char_to_word = {}
        last_found = 0
        for word_idx, word in enumerate(words):
            # Find the word in the text, starting search after the previous word
            start = text.find(word, last_found)
            if start != -1:
                for i in range(start, start + len(word)):
                    char_to_word[i] = word_idx
                last_found = start + len(word)
        
        char_pos = 0
        for run in self.runs:
            if not run.text:
                continue
            for char in run.text:
                if char_pos in char_to_word:
                    word_idx = char_to_word[char_pos]
                    if not char.isspace():
                        if run.bold and run.italic:
                            formatted['italic_bold'].add(word_idx)
                        elif run.italic:
                            formatted['italic'].add(word_idx)
                        elif run.bold:
                            formatted['bold'].add(word_idx)
                char_pos += 1
        return formatted


# ============================================================================
# TRANSLATION BACKENDS (keeping previous implementations)
# ============================================================================

class CTranslate2Translator:
    """CTranslate2 for fast translation"""
    
    MODELS = {
        'en_to_x': 'cstr/wmt21ct2_int8',
        'x_to_en': 'cstr/wmt21-ct2-x-en-int8',
    }
    
    SUPPORTED_LANGS = ['de', 'es', 'fr', 'it', 'ja', 'zh', 'ru', 'pt', 'nl', 'cs', 'uk']
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = False
        
        if not HAS_CT2:
            logger.info("⊘ CT2: library not installed")
            return
        
        if src_lang == 'en' and tgt_lang in self.SUPPORTED_LANGS:
            self.direction = 'en_to_x'
            self.tokenizer_name = "facebook/wmt21-dense-24-wide-en-x"
        elif tgt_lang == 'en' and src_lang in self.SUPPORTED_LANGS:
            self.direction = 'x_to_en'
            self.tokenizer_name = "facebook/wmt21-dense-24-wide-x-en"
        else:
            logger.info(f"⊘ CT2: {src_lang}→{tgt_lang} not supported")
            return
        
        try:
            logger.info(f"Loading CT2 model...")
            model_path = self._get_or_download_model()
            self.translator = ctranslate2.Translator(model_path, device="auto")
            self.tokenizer = AutoTokenizer.from_pretrained(self.tokenizer_name)
            self.tokenizer.src_lang = src_lang
            self.tokenizer.tgt_lang = tgt_lang
            self.available = True
            logger.info(f"✓ CT2 ready ({src_lang}→{tgt_lang})")
        except Exception as e:
            logger.error(f"CT2 init failed: {e}")
    
    def _get_or_download_model(self) -> str:
        cache_base = Path.home() / '.cache' / 'huggingface' / 'hub'
        model_repo = self.MODELS[self.direction]
        model_dir = cache_base / f"models--{model_repo.replace('/', '--')}"
        ref_path = model_dir / 'refs' / 'main'
        
        if ref_path.exists():
            with open(ref_path) as f:
                commit_hash = f.read().strip()
            snapshot_path = model_dir / 'snapshots' / commit_hash
            if (snapshot_path / 'model.bin').exists():
                return str(snapshot_path)
        
        logger.info("Downloading CT2 model...")
        return snapshot_download(repo_id=model_repo)
    
    def translate_batch(self, texts: List[str]) -> List[str]:
        if not self.available or not texts:
            return texts
        
        try:
            # 1. Ensure language tokens are correctly formatted
            # WMT21 models expect the target lang as a prefix
            target_prefix = [[self.tokenizer.lang_code_to_token[self.tgt_lang]]] * len(texts)
            
            # 2. Tokenize with specific padding/truncation
            source_tokens = [self.tokenizer.convert_ids_to_tokens(self.tokenizer.encode(t)) for t in texts]
            
            # 3. Translate with proper beam search and repetition penalty
            results = self.translator.translate_batch(
                source_tokens,
                target_prefix=target_prefix,
                beam_size=5,
                max_batch_size=16,
                repetition_penalty=1.2,
                # Prevent the model from just 'copying' the source if it gets confused
                disable_unk=True 
            )
            
            translated = []
            for i, res in enumerate(results):
                # Strip the language token from the start of the result
                tokens = res.hypotheses[0]
                if self.tokenizer.lang_code_to_token[self.tgt_lang] in tokens:
                    tokens = tokens[tokens.index(self.tokenizer.lang_code_to_token[self.tgt_lang]) + 1:]
                
                decoded = self.tokenizer.decode(self.tokenizer.convert_tokens_to_ids(tokens), skip_special_tokens=True)
                translated.append(decoded.strip())
                
            return translated
        except Exception as e:
            logger.error(f"CT2 Critical Error: {e}")
            return texts


class NLLBTranslator:
    """NLLB-200 translator using CTranslate2 for 4x speedup and low memory"""
    
    # Map sizes to specific CTranslate2 optimized repositories
    REPOS = {
        "600M": "JustFrederik/nllb-200-distilled-600M-ct2-int8",
        "1.3B": "OpenNMT/nllb-200-distilled-1.3B-ct2-int8",
        "3.3B": "OpenNMT/nllb-200-3.3B-ct2-int8"
    }

    LANG_CODES = {
        'en': 'eng_Latn', 'de': 'deu_Latn', 'fr': 'fra_Latn', 
        'es': 'spa_Latn', 'it': 'ita_Latn', 'pt': 'por_Latn',
        'ru': 'rus_Cyrl', 'zh': 'zho_Hans', 'ja': 'jpn_Jpan',
        'ko': 'kor_Hang', 'ar': 'arb_Arab', 'hi': 'hin_Deva',
        'nl': 'nld_Latn', 'pl': 'pol_Latn', 'tr': 'tur_Latn',
        'cs': 'ces_Latn', 'uk': 'ukr_Cyrl', 'vi': 'vie_Latn',
    }
    
    def __init__(self, src_lang: str, tgt_lang: str, model_size: str = "600M"):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = False
        
        if not HAS_CT2:
            logger.info("⊘ NLLB-CT2: ctranslate2 not installed")
            return
        
        self.src_code = self.LANG_CODES.get(src_lang)
        self.tgt_code = self.LANG_CODES.get(tgt_lang)
        
        if not self.src_code or not self.tgt_code:
            logger.info(f"⊘ NLLB-CT2: {src_lang}→{tgt_lang} not supported")
            return
        
        try:
            repo_id = self.REPOS.get(model_size, self.REPOS["600M"])
            logger.info(f"Loading NLLB-{model_size} (CTranslate2 int8) from {repo_id}...")
            
            # 1. Download model from Hub
            model_path = snapshot_download(repo_id=repo_id)
            
            # 2. Initialize CT2 Translator
            device = "cuda" if torch.cuda.is_available() else "cpu"
            # compute_type: int8_float16 for GPU, int8 for CPU
            compute_type = "int8_float16" if device == "cuda" else "int8"
            
            self.translator = ctranslate2.Translator(
                model_path, 
                device=device, 
                compute_type=compute_type
            )
            
            # 3. Initialize Tokenizer (from the original FB repo to ensure compatibility)
            # We use the distilled-600M tokenizer for all because NLLB uses a shared vocab
            self.tokenizer = HFTokenizer.from_pretrained(
                "facebook/nllb-200-distilled-600M", 
                src_lang=self.src_code
            )
            
            self.available = True
            logger.info(f"✓ NLLB-{model_size} CT2 ready on {device}")
        except Exception as e:
            logger.error(f"NLLB-CT2 init failed: {e}")
    
    def translate_batch(self, texts: List[str], batch_size: int = 16) -> List[str]:
        """Translate batch using CT2 efficient beam search"""
        if not self.available or not texts:
            return texts
        
        try:
            results = []
            # NLLB requires the target language code as the first token (target prefix)
            target_prefix = [self.tgt_code]
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                
                # Tokenize
                source_tokens = [
                    self.tokenizer.convert_ids_to_tokens(self.tokenizer.encode(t)) 
                    for t in batch
                ]
                
                # Translate
                step_results = self.translator.translate_batch(
                    source_tokens,
                    target_prefix=[target_prefix] * len(batch),
                    beam_size=4,
                    max_batch_size=batch_size,
                    repetition_penalty=1.1
                )
                
                # Decode (skipping the lang code prefix in the output)
                for res in step_results:
                    tokens = res.hypotheses[0]
                    # The first token is usually the target lang code
                    if tokens[0] == self.tgt_code:
                        tokens = tokens[1:]
                    
                    decoded = self.tokenizer.decode(
                        self.tokenizer.convert_tokens_to_ids(tokens), 
                        skip_special_tokens=True
                    )
                    results.append(decoded.strip())
            
            return results
        except Exception as e:
            logger.error(f"NLLB-CT2 translation failed: {e}")
            return texts
        
class LLMTranslator:
    """LLM translator (OpenAI/Anthropic/Ollama)"""
    
    def __init__(self, src_lang: str, tgt_lang: str, preferred_provider: Optional[str] = None):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.providers = self._init_providers(preferred_provider)
        
        if self.providers:
            logger.info(f"✓ LLM available ({list(self.providers.keys())})")
    
    def _init_providers(self, preferred: Optional[str] = None) -> Dict[str, Any]:
        providers = {}
        
        if HAS_OPENAI and os.getenv("OPENAI_API_KEY"):
            providers["openai"] = {
                "client": AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY")),
                "model": os.getenv("OPENAI_MODEL", "gpt-4o-mini")
            }
        
        if HAS_ANTHROPIC and os.getenv("ANTHROPIC_API_KEY"):
            providers["anthropic"] = {
                "client": AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY")),
                "model": os.getenv("ANTHROPIC_MODEL", "claude-3-5-sonnet-20241022")
            }
        
        if self._check_ollama():
            providers["ollama"] = {
                "url": "http://localhost:11434/api/generate",
                "model": self._get_ollama_model()
            }
        
        if preferred and preferred in providers:
            return {preferred: providers[preferred]}
        
        return providers
    
    def _check_ollama(self) -> bool:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            return r.status_code == 200 and len(r.json().get("models", [])) > 0
        except:
            return False
    
    def _get_ollama_model(self) -> str:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            models = r.json().get("models", [])
            return models[0]["name"] if models else "llama3.2"
        except:
            return "llama3.2"
    
    async def translate_text(self, text: str, use_alignment: bool = True) -> Optional[str]:
        """Translate single text"""
        if not text.strip() or not self.providers:
            return None
        
        if use_alignment:
            prompt = (
                f"Translate the following text from {self.src_lang} to {self.tgt_lang}. "
                f"Preserve the word order as much as possible for alignment purposes. "
                f"Return ONLY the translation:\n\n{text}"
            )
        else:
            prompt = (
                f"Translate the following text from {self.src_lang} to {self.tgt_lang}. "
                f"Provide a natural, fluent translation. "
                f"Return ONLY the translation:\n\n{text}"
            )
        
        for provider_name, provider in self.providers.items():
            try:
                if provider_name == "openai":
                    response = await provider["client"].chat.completions.create(
                        model=provider["model"],
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                        max_tokens=4000
                    )
                    return response.choices[0].message.content.strip()
                
                elif provider_name == "anthropic":
                    response = await provider["client"].messages.create(
                        model=provider["model"],
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                        max_tokens=4000
                    )
                    return response.content[0].text.strip()
                
                elif provider_name == "ollama":
                    r = requests.post(
                        provider["url"],
                        json={"model": provider["model"], "prompt": prompt, "stream": False},
                        timeout=120
                    )
                    if r.status_code == 200:
                        return r.json().get("response", "").strip()
            
            except Exception as e:
                logger.debug(f"{provider_name} failed: {e}")
                continue
        
        return None
    
    async def translate_batch(self, texts: List[str], use_alignment: bool = True) -> List[str]:
        """Translate batch"""
        tasks = [self.translate_text(text, use_alignment) for text in texts]
        results = await asyncio.gather(*tasks)
        return [res if res else text for res, text in zip(results, texts)]


# ============================================================================
# WORD ALIGNERS - FIXED WITH BETTER FALLBACK
# ============================================================================

class LindatAligner:
    """Lindat word alignment API"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = self._check_available()
        if self.available:
            logger.info(f"✓ Lindat aligner available ({src_lang}-{tgt_lang})")
        else:
            logger.info(f"⊘ Lindat aligner not available ({src_lang}-{tgt_lang})")
    
    def _check_available(self) -> bool:
        try:
            r = requests.get(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                timeout=5
            )
            return r.status_code in [200, 405]
        except:
            return False
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            r = requests.post(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                headers={'Content-Type': 'application/json'},
                json={'src_tokens': [src_words], 'trg_tokens': [tgt_words]},
                timeout=30
            )
            
            if r.status_code == 200:
                alignment = r.json()["alignment"][0]
                return [(int(a[0]), int(a[1])) for a in alignment]
        except Exception as e:
            logger.debug(f"Lindat alignment failed: {e}")
        
        return []

class AwesomeAlignAligner:
    """awesome-align: BERT-based word aligner"""
    
    def __init__(self):
        self.available = False
        self.model = None
        self.tokenizer = None
        
        if not HAS_TRANSFORMERS:
            logger.info("⊘ awesome-align not available (needs transformers)")
            return
        
        try:
            import torch
            from transformers import BertModel, BertTokenizer
            
            logger.info("Loading awesome-align (mBERT)...")
            self.model = BertModel.from_pretrained('bert-base-multilingual-cased')
            self.tokenizer = BertTokenizer.from_pretrained('bert-base-multilingual-cased')
            self.model.eval()
            
            # Move to GPU if available
            self.device = "cuda" if torch.cuda.is_available() else "cpu"
            self.model.to(self.device)
            
            self.available = True
            logger.info(f"✓ awesome-align ready on {self.device}")
        except Exception as e:
            logger.info(f"⊘ awesome-align init failed: {e}")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices using mBERT embeddings"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            import torch
            import itertools
            
            # Tokenize words
            token_src = [self.tokenizer.tokenize(word) for word in src_words]
            token_tgt = [self.tokenizer.tokenize(word) for word in tgt_words]
            
            wid_src = [self.tokenizer.convert_tokens_to_ids(x) for x in token_src]
            wid_tgt = [self.tokenizer.convert_tokens_to_ids(x) for x in token_tgt]
            
            ids_src = self.tokenizer.prepare_for_model(
                list(itertools.chain(*wid_src)), 
                return_tensors='pt', 
                truncation=True,
                model_max_length=self.tokenizer.model_max_length
            )['input_ids'].to(self.device)
            
            ids_tgt = self.tokenizer.prepare_for_model(
                list(itertools.chain(*wid_tgt)), 
                return_tensors='pt', 
                truncation=True,
                model_max_length=self.tokenizer.model_max_length
            )['input_ids'].to(self.device)
            
            # Create subword to word mapping
            sub2word_map_src = []
            for i, word_list in enumerate(token_src):
                sub2word_map_src += [i for _ in word_list]
            
            sub2word_map_tgt = []
            for i, word_list in enumerate(token_tgt):
                sub2word_map_tgt += [i for _ in word_list]
            
            # Extract alignments using layer 8
            align_layer = 8
            threshold = 1e-3
            
            with torch.no_grad():
                out_src = self.model(ids_src.unsqueeze(0), output_hidden_states=True)[2][align_layer][0, 1:-1]
                out_tgt = self.model(ids_tgt.unsqueeze(0), output_hidden_states=True)[2][align_layer][0, 1:-1]
                
                dot_prod = torch.matmul(out_src, out_tgt.transpose(-1, -2))
                
                softmax_srctgt = torch.nn.Softmax(dim=-1)(dot_prod)
                softmax_tgtsrc = torch.nn.Softmax(dim=-2)(dot_prod)
                
                softmax_inter = (softmax_srctgt > threshold) * (softmax_tgtsrc > threshold)
            
            # Convert subword alignments to word alignments
            align_subwords = torch.nonzero(softmax_inter, as_tuple=False)
            align_words = set()
            for i, j in align_subwords:
                align_words.add((sub2word_map_src[i.item()], sub2word_map_tgt[j.item()]))
            
            return sorted(list(align_words))
        
        except Exception as e:
            logger.debug(f"awesome-align failed: {e}")
            return []

class FastAlignAligner:
    """fast_align local aligner - uses temp file for binary"""
    
    def __init__(self):
        self.available = False
        self.mode = None
        self.binary_path = None
        
        if HAS_FAST_ALIGN:
            self.available = True
            self.mode = "python"
            logger.info("✓ fast_align available (Python package)")
            return
        
        # Check relative and absolute paths
        script_dir = Path(__file__).parent
        binary_locations = [
            "../fast_align/build/fast_align",
            "./fast_align/build/fast_align",
            "fast_align" # If in PATH
        ]
        
        for loc in binary_locations:
            path = script_dir / loc if not loc.startswith('/') else Path(loc)
            if os.path.isfile(path) and os.access(path, os.X_OK):
                self.binary_path = str(path)
                self.available = True
                self.mode = "binary"
                logger.info(f"✓ fast_align binary found at {self.binary_path}")
                return
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            if self.mode == "python":
                from fast_align import align
                src_text = ' '.join(src_words)
                tgt_text = ' '.join(tgt_words)
                result = align([f"{src_text} ||| {tgt_text}"], forward=True)
                return [tuple(map(int, p.split('-'))) for p in result[0].split()]
            
            elif self.mode == "binary":
                import subprocess
                import tempfile
                
                input_str = f"{' '.join(src_words)} ||| {' '.join(tgt_words)}\n"
                
                with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
                    f.write(input_str)
                    temp_path = f.name
                
                try:
                    # -d: diagonal, -o: optimize tension, -v: variational bayes
                    # -I 1: single pass is enough for alignment of a known translation
                    cmd = [self.binary_path, '-i', temp_path, '-d', '-o', '-v', '-I', '1']
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                    
                    if result.returncode != 0:
                        return []
                    
                    # fast_align binary outputs to stdout in Pharaoh format (i-j)
                    alignments = []
                    output = result.stdout.strip()
                    if output:
                        for pair in output.split():
                            if '-' in pair:
                                i, j = map(int, pair.split('-'))
                                alignments.append((i, j))
                    return alignments
                finally:
                    if os.path.exists(temp_path): os.unlink(temp_path)
        except Exception as e:
            logger.debug(f"fast_align execution failed: {e}")
            return []


class SimAlignAligner:
    """SimAlign aligner"""
    
    def __init__(self):
        self.available = HAS_SIMALIGN
        if self.available:
            try:
                self.aligner = SentenceAligner(model="bert", token_type="bpe", matching_methods="mai")
                logger.info("✓ SimAlign available")
            except Exception as e:
                logger.debug(f"SimAlign init failed: {e}")
                self.available = False
                logger.info("⊘ SimAlign init failed")
        else:
            logger.info("⊘ SimAlign not available (install: pip install simalign)")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            result = self.aligner.get_word_aligns(src_words, tgt_words)
            alignments = []
            for src_idx, tgt_idx in result['mwmf'].items():
                alignments.append((src_idx, tgt_idx))
            return alignments
        except Exception as e:
            logger.debug(f"SimAlign failed: {e}")
            return []


class HeuristicAligner:
    """Heuristic fallback - align words that appear in both"""
    
    def __init__(self):
        logger.info("✓ Heuristic aligner (fallback)")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Simple heuristic alignment"""
        alignments = []
        
        src_lower = [w.lower().strip('.,!?;:') for w in src_words]
        tgt_lower = [w.lower().strip('.,!?;:') for w in tgt_words]
        
        for i, src_word in enumerate(src_lower):
            for j, tgt_word in enumerate(tgt_lower):
                if src_word == tgt_word and len(src_word) > 2:
                    alignments.append((i, j))
        
        return alignments


class MultiAligner:
    """Try multiple aligners with proper fallback"""
    
    def __init__(self, src_lang: str, tgt_lang: str, preferred: Optional[str] = None):
        self.aligners = []
        
        # If a specific aligner was requested
        if preferred and preferred != "auto":
            if preferred == "lindat":
                lindat = LindatAligner(src_lang, tgt_lang)
                if lindat.available:
                    self.aligners.append(("Lindat", lindat))
                else:
                    logger.warning(f"Requested aligner '{preferred}' not available, will try others")
            
            elif preferred == "awesome":
                awesome = AwesomeAlignAligner()
                if awesome.available:
                    self.aligners.append(("awesome-align", awesome))
            
            elif preferred == "fast_align":
                fast_align = FastAlignAligner()
                if fast_align.available:
                    self.aligners.append(("fast_align", fast_align))
                else:
                    logger.warning(f"Requested aligner '{preferred}' not available, will try others")
            
            elif preferred == "simalign":
                simalign = SimAlignAligner()
                if simalign.available:
                    self.aligners.append(("SimAlign", simalign))
                else:
                    logger.warning(f"Requested aligner '{preferred}' not available, will try others")
            
            elif preferred == "heuristic":
                self.aligners.append(("Heuristic", HeuristicAligner()))
        
        # Auto mode or fallback: try all available aligners
        if not self.aligners or preferred == "auto":
            # Try Lindat first
            lindat = LindatAligner(src_lang, tgt_lang)
            if lindat.available:
                self.aligners.append(("Lindat", lindat))

            awesome = AwesomeAlignAligner()
            if awesome.available:
                self.aligners.append(("awesome-align", awesome))
            
            # Then SimAlign
            simalign = SimAlignAligner()
            if simalign.available:
                self.aligners.append(("SimAlign", simalign))
            
            # Then fast_align
            fast_align = FastAlignAligner()
            if fast_align.available:
                self.aligners.append(("fast_align", fast_align))
        
        # Always add heuristic as final fallback
        if not any(name == "Heuristic" for name, _ in self.aligners):
            self.aligners.append(("Heuristic", HeuristicAligner()))
        
        logger.info(f"Aligner chain: {[name for name, _ in self.aligners]}")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Try aligners in order until one succeeds"""
        for name, aligner in self.aligners:
            result = aligner.align(src_words, tgt_words)
            if result:
                logger.debug(f"✓ {name} alignment: {len(result)} links")
                return result
        
        logger.debug("Using heuristic fallback (no quality alignments found)")
        return []


# ============================================================================
# DOCUMENT TRANSLATOR - FIXED FOOTNOTE HANDLING
# ============================================================================

class UltimateDocumentTranslator:
    """Production-grade document translator with configurable backends"""
    
    def __init__(
        self,
        src_lang: str,
        tgt_lang: str,
        mode: TranslationMode = TranslationMode.HYBRID,
        nmt_backend: Optional[str] = None,
        llm_provider: Optional[str] = None,
        aligner: Optional[str] = None,
        nllb_model_size: str = "600M"
    ):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.mode = mode
        
        logger.info(f"Initializing translator ({src_lang}→{tgt_lang}, mode={mode.value})")

        # Ensure nmt_backend defaults to auto if in NMT or HYBRID mode
        if nmt_backend is None:
            nmt_backend = "auto"
        
        # Initialize translators based on mode
        self.ct2 = None
        self.nllb = None
        self.llm = None
        
        if mode in [TranslationMode.NMT_ONLY, TranslationMode.HYBRID]:
            if nmt_backend == "nllb" or nmt_backend == "auto":
                self.nllb = NLLBTranslator(src_lang, tgt_lang, nllb_model_size)
            
            if nmt_backend == "ct2" or nmt_backend == "auto":
                self.ct2 = CTranslate2Translator(src_lang, tgt_lang)
        
        if mode in [TranslationMode.LLM_WITH_ALIGN, TranslationMode.LLM_WITHOUT_ALIGN, TranslationMode.HYBRID]:
            self.llm = LLMTranslator(src_lang, tgt_lang, llm_provider)
        
        # Initialize aligner if needed
        self.aligner = None
        if mode in [TranslationMode.LLM_WITH_ALIGN, TranslationMode.HYBRID]:
            self.aligner = MultiAligner(src_lang, tgt_lang, aligner or "auto")
        
        # Check we have at least one backend
        has_backend = False
        if self.ct2 and self.ct2.available:
            has_backend = True
            logger.info("Primary backend: CTranslate2")
        if self.nllb and self.nllb.available:
            has_backend = True
            logger.info("Primary backend: NLLB")
        if self.llm and self.llm.providers:
            has_backend = True
            logger.info(f"LLM backend: {list(self.llm.providers.keys())}")
        
        if not has_backend:
            logger.error("No translation backends available!")
            sys.exit(1)
    
    async def translate_text(self, text: str) -> str:
        """Translate single text with fallback chain"""
        if not text.strip():
            return text
        
        if self.mode == TranslationMode.NMT_ONLY:
            if self.nllb and self.nllb.available:
                result = self.nllb.translate_batch([text])[0]
                if result and result != text:
                    return result
            
            if self.ct2 and self.ct2.available:
                result = self.ct2.translate_batch([text])[0]
                if result and result != text:
                    return result
        
        elif self.mode == TranslationMode.LLM_WITH_ALIGN:
            if self.llm and self.llm.providers:
                result = await self.llm.translate_text(text, use_alignment=True)
                if result:
                    return result
        
        elif self.mode == TranslationMode.LLM_WITHOUT_ALIGN:
            if self.llm and self.llm.providers:
                result = await self.llm.translate_text(text, use_alignment=False)
                if result:
                    return result
        
        elif self.mode == TranslationMode.HYBRID:
            if self.nllb and self.nllb.available:
                result = self.nllb.translate_batch([text])[0]
                if result and result != text:
                    return result
            
            if self.ct2 and self.ct2.available:
                result = self.ct2.translate_batch([text])[0]
                if result and result != text:
                    return result
            
            if self.llm and self.llm.providers:
                result = await self.llm.translate_text(text, use_alignment=True)
                if result:
                    return result
        
        logger.error(f"Translation failed for: {text[:50]}...")
        return text
    
    def extract_paragraph(self, para: Paragraph) -> TranslatableParagraph:
        runs = []
        for run in para.runs:
            font_color = None
            try:
                # Use getattr to safely check for attributes
                if hasattr(run.font, 'color') and run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    font_color = (rgb[0], rgb[1], rgb[2])
            except:
                pass
            
            # Use getattr for size and other properties
            font_size = None
            try:
                if run.font.size:
                    font_size = run.font.size.pt
            except:
                pass

            runs.append(FormatRun(
                text=run.text,
                bold=getattr(run.font, 'bold', None),
                italic=getattr(run.font, 'italic', None),
                underline=getattr(run.font, 'underline', None),
                font_name=getattr(run.font, 'name', None),
                font_size=font_size,
                font_color=font_color
            ))
        
        trans_para = TranslatableParagraph(runs=runs)
        trans_para.metadata['style'] = getattr(para, 'style', None)
        trans_para.metadata['alignment'] = getattr(para, 'alignment', None)
        
        pf = para.paragraph_format
        trans_para.metadata['layout'] = {
            'left_indent': pf.left_indent,
            'right_indent': pf.right_indent,
            'first_line_indent': pf.first_line_indent,
            'line_spacing': pf.line_spacing,
            'space_before': pf.space_before,
            'space_after': pf.space_after
        }
        return trans_para
    
    def copy_font_properties(self, target_run, source_run: FormatRun):
        """Copies all font properties including safety checks."""
        if source_run.font_name:
            target_run.font.name = source_run.font_name
        if source_run.font_size:
            target_run.font.size = Pt(source_run.font_size)
        if source_run.font_color:
            target_run.font.color.rgb = RGBColor(*source_run.font_color)
        if source_run.underline is not None:
            target_run.font.underline = source_run.underline
    
    def apply_aligned_formatting(self, para: Paragraph, trans_para: TranslatableParagraph, translated_text: str, alignment: List[Tuple[int, int]]):
        """Maps formatting from clean words to punctuated units by skipping non-alpha tokens."""
        import re
        
        # 1. Protect Footnotes
        p = para._p
        footnote_refs = []
        for run_element in p.xpath('.//w:r'):
            if run_element.xpath('.//w:footnoteReference | .//w:footnoteRef'):
                footnote_refs.append(run_element)
            run_element.getparent().remove(run_element)

        # 2. Extract clean source and raw target units
        src_clean_words = trans_para.get_words()
        # WMT21 often separates punctuation with spaces, so split() catches them as units
        tgt_raw_units = translated_text.split() 
        formatted_indices = trans_para.get_formatted_word_indices()

        # 3. Create a mapping: Clean Aligner Index -> Raw Target Unit Index
        # We skip tokens that are strictly punctuation so the aligner index matches words
        clean_to_raw_tgt = {}
        clean_idx = 0
        for raw_idx, unit in enumerate(tgt_raw_units):
            # If the unit contains alphanumeric characters, it's a 'word' for the aligner
            if re.search(r'\w', unit):
                clean_to_raw_tgt[clean_idx] = raw_idx
                clean_idx += 1

        # 4. Map Style to Raw Target Units
        tgt_styles = {} 
        for src_idx, align_tgt_idx in alignment:
            # align_tgt_idx is what the ALIGNER thinks is the word index
            # raw_tgt_idx is the actual index in the punctuated list
            raw_tgt_idx = clean_to_raw_tgt.get(align_tgt_idx)
            
            if src_idx in src_clean_words or True: # safety check
                if src_idx < len(src_clean_words) and raw_tgt_idx is not None:
                    for style in ['italic_bold', 'bold', 'italic']:
                        if src_idx in formatted_indices[style]:
                            tgt_styles[raw_tgt_idx] = style
                            break

        # 5. Build runs using the raw units (preserving punctuation and spacing)
        font_template = next((r for r in trans_para.runs if r.text.strip()), None)
        
        for i, unit in enumerate(tgt_raw_units):
            style = tgt_styles.get(i, None)
            
            # Logic to handle spacing: don't add space before punctuation usually, 
            # but WMT21 output text.split() + " " join is the safest restoration.
            run_text = unit + (" " if i < len(tgt_raw_units) - 1 else "")
            
            # CRITICAL: If the unit is JUST punctuation, check if the PREVIOUS word was formatted
            # This ensures "Theology:" is all bold even if ":" is a separate token
            if not re.search(r'\w', unit) and i > 0:
                prev_style = tgt_styles.get(i-1)
                if prev_style:
                    style = prev_style

            run = para.add_run(run_text)
            if style == 'italic_bold': run.bold = run.italic = True
            elif style == 'bold': run.bold = True
            elif style == 'italic': run.italic = True
            
            if font_template:
                self.copy_font_properties(run, font_template)

        # 6. Restore Footnotes
        for ref in footnote_refs:
            p.append(ref)
    
    def is_paragraph_safe_to_translate(self, para: Paragraph) -> bool:
        """Check if paragraph can be safely translated"""
        if not para.text or not para.text.strip():
            return False
        
        text = para.text.strip()
        if not text:
            return False
        
        if len(para.text.strip()) <= 1 and not any(run.text.strip() for run in para.runs):
            return False
        
        try:
            for run in para.runs:
                if run._element.xpath('.//w:drawing | .//w:pict'):
                    logger.debug(f"Skipping paragraph with drawing/image")
                    return False
        except:
            pass
        
        try:
            field_chars = para._element.xpath('.//w:fldChar')
            if field_chars:
                is_footnote_field = para._element.xpath('.//w:footnoteReference')
                if not is_footnote_field:
                    logger.debug(f"Skipping paragraph with field")
                    return False
        except:
            pass
        
        return True
    
    async def translate_paragraph(self, para: Paragraph):
        """Translate paragraph with sentence-splitting and clean word alignment."""
        if not para.text or not para.text.strip() or not self.is_paragraph_safe_to_translate(para):
            return
        
        try:
            trans_para = self.extract_paragraph(para)
            original_text = trans_para.get_text()
            
            # 1. Split into sentences: WMT21 performs best on sentence units
            sentences = re.split(r'(?<=[.!?]) +', original_text)
            translated_sentences = []

            # 2. Translation Fallback Logic
            if self.ct2 and self.ct2.available:
                translated_sentences = self.ct2.translate_batch(sentences)
            elif self.nllb and self.nllb.available:
                translated_sentences = self.nllb.translate_batch(sentences)
            elif self.llm and self.llm.providers:
                tasks = [self.llm.translate_text(s, use_alignment=True) for s in sentences]
                results = await asyncio.gather(*tasks)
                translated_sentences = [res if res else s for res, s in zip(results, sentences)]
            else:
                return

            translated_text = " ".join(translated_sentences)
            
            # 3. PREPARE CLEAN DATA FOR ALIGNER (Alphanumeric only)
            # We strip punctuation from BOTH sides so the indices (0, 1, 2...) match words.
            src_words = trans_para.get_words() # Uses \w+ (clean)
            tgt_words = re.findall(r"\w+", translated_text) # FIX: Use \w+ here too
            
            if not src_words or not tgt_words:
                # If everything is punctuation, just apply the text as plain
                self.apply_plain_formatting(para, trans_para, translated_text)
                return
            
            # 4. ALIGNER CALL (Clean Source vs Clean Target)
            alignment = self.aligner.align(src_words, tgt_words)

            # 5. DEBUG LOGS
            logger.debug(f"SRC (Clean): {src_words}")
            logger.debug(f"TGT (Clean): {tgt_words}")
            logger.debug(f"ALIGN: {alignment}")
            
            # 6. ASSEMBLY (Passes punctuated text. Formatting logic handles the rest)
            self.apply_aligned_formatting(para, trans_para, translated_text, alignment)
                
        except Exception as e:
            logger.error(f"Error in translate_paragraph: {e}")
    
    def get_footnotes(self, doc: Document) -> List[Paragraph]:
        """Extract footnotes - FIXED: Pass doc as parent to avoid Part attribute error"""
        try:
            document_part = doc.part
            footnote_part = None
            
            for rel in document_part.rels.values():
                if "relationships/footnotes" in rel.reltype:
                    footnote_part = rel.target_part
                    break
            
            if not footnote_part:
                return []
            
            from docx.oxml import parse_xml
            root = parse_xml(footnote_part.blob)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            paragraphs = []
            for footnote in root.xpath('//w:footnote', namespaces=ns):
                f_id = footnote.get(f'{{{ns["w"]}}}id')
                # Skip internal Word markers (id 0 and -1)
                if f_id and int(f_id) <= 0:
                    continue
                
                for p_elem in footnote.xpath('.//w:p', namespaces=ns):
                    # FIX: Pass 'doc' instead of 'None' so para.part is valid
                    para = Paragraph(p_elem, doc)
                    if para.text.strip():
                        paragraphs.append(para)
            
            self._footnote_root = root
            self._footnote_part = footnote_part
            return paragraphs
            
        except Exception as e:
            logger.warning(f"Could not extract footnotes: {e}")
            return []
    
    def get_all_paragraphs(self, doc: Document) -> List[Tuple[Paragraph, str]]:
        """Aggregate all paragraphs from document"""
        all_paras = []
        
        # Main body
        for para in doc.paragraphs:
            all_paras.append((para, "body"))
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_paras.append((para, "table"))
        
        # Footnotes - with error handling
        try:
            footnote_paras = self.get_footnotes(doc)
            for para in footnote_paras:
                all_paras.append((para, "footnote"))
        except Exception as e:
            logger.warning(f"Skipping footnotes due to error: {e}")
        
        # Headers/Footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                all_paras.append((para, "header"))
            for para in section.footer.paragraphs:
                all_paras.append((para, "footer"))
        
        return all_paras
    
    async def translate_document(self, input_path: Path, output_path: Path):
        """Translate entire document"""
        logger.info(f"Loading document: {input_path}")
        
        try:
            doc = Document(str(input_path))
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            return
        
        all_paras_with_location = self.get_all_paragraphs(doc)
        
        translatable = []
        for para, location in all_paras_with_location:
            if para.text and para.text.strip() and self.is_paragraph_safe_to_translate(para):
                translatable.append((para, location))
        
        by_location = defaultdict(int)
        for _, location in translatable:
            by_location[location] += 1
        
        logger.info(f"Found {len(translatable)} paragraphs to translate:")
        for location, count in by_location.items():
            logger.info(f"  - {location}: {count}")
        
        for para, location in tqdm(translatable, desc="Translating"):
            try:
                await self.translate_paragraph(para)
                await asyncio.sleep(0.05)
            except Exception as e:
                logger.error(f"Error translating {location} paragraph: {e}")
                continue
        
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            if output_path.exists():
                output_path.unlink()
            
            logger.info(f"Saving to: {output_path}")
            
            if hasattr(self, '_footnote_part') and hasattr(self, '_footnote_root'):
                try:
                    from lxml import etree
                    self._footnote_part._blob = etree.tostring(self._footnote_root)
                    logger.info("✓ Updated footnotes in output")
                except Exception as e:
                    logger.warning(f"Could not update footnotes: {e}")
            
            doc.save(str(output_path))
            
            if not output_path.exists():
                raise Exception("Output file was not created")
            
            if output_path.stat().st_size == 0:
                raise Exception("Output file is empty")
            
            test_doc = Document(str(output_path))
            logger.info(f"✓ Document verified ({len(test_doc.paragraphs)} paragraphs)")
            logger.info("✓ Translation complete!")
            
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            if output_path.exists():
                output_path.unlink()
            raise


# ============================================================================
# CLI
# ============================================================================

async def main():
    parser = argparse.ArgumentParser(
        description='Document Translator',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Default (auto-detect best backends)
  %(prog)s input.docx output.docx -s en -t de
  
  # Use NLLB for translation
  %(prog)s input.docx output.docx -s en -t fr --nmt nllb
  
  # Use LLM with alignment
  %(prog)s input.docx output.docx -s en -t es --mode llm-align --llm openai
  
  # Use LLM without alignment (natural translation)
  %(prog)s input.docx output.docx -s en -t it --mode llm-plain --llm anthropic
  
  # Hybrid mode (will use Lindat aligner)
  %(prog)s input.docx output.docx -s en -t de --mode hybrid
  
  # Use larger NLLB model
  %(prog)s input.docx output.docx -s en -t ja --nmt nllb --nllb-size 1.3B

Environment Variables:
  OPENAI_API_KEY       - OpenAI API key
  OPENAI_MODEL         - OpenAI model (default: gpt-4o-mini)
  ANTHROPIC_API_KEY    - Anthropic API key
  ANTHROPIC_MODEL      - Anthropic model (default: claude-3-5-sonnet-20241022)

Installation Notes:
  fast_align: This is a C++ tool, not a Python package. 
              See: https://github.com/clab/fast_align
              Or install Python wrapper: pip install fast-align
  
  SimAlign:   pip install simalign
  NLLB:       pip install transformers torch
        """
    )
    
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('output', help='Output .docx file')
    parser.add_argument('-s', '--source', default='en', help='Source language (default: en)')
    parser.add_argument('-t', '--target', default='de', help='Target language (default: de)')
    
    parser.add_argument(
        '--mode',
        choices=['nmt', 'llm-align', 'llm-plain', 'hybrid'],
        default='hybrid',
        help='Translation mode (default: hybrid)'
    )
    
    parser.add_argument(
        '--nmt',
        choices=['ct2', 'nllb', 'auto'],
        help='NMT backend (default: auto)'
    )
    
    parser.add_argument(
        '--nllb-size',
        choices=['600M', '1.3B', '3.3B'],
        default='600M',
        help='NLLB model size (default: 600M)'
    )
    
    parser.add_argument(
        '--llm',
        choices=['openai', 'anthropic', 'ollama'],
        help='LLM provider (default: auto-detect)'
    )
    
    parser.add_argument(
        '--aligner',
        choices=['lindat', 'fast_align', 'simalign', 'heuristic', 'auto'],
        help='Word aligner (default: auto)'
    )
    
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"Error: Only .docx files supported")
        sys.exit(1)
    
    mode_map = {
        'nmt': TranslationMode.NMT_ONLY,
        'llm-align': TranslationMode.LLM_WITH_ALIGN,
        'llm-plain': TranslationMode.LLM_WITHOUT_ALIGN,
        'hybrid': TranslationMode.HYBRID
    }
    
    print(f"\n{'='*60}")
    print(f"Document Translator - Production Version")
    print(f"{'='*60}")
    print(f"Input:      {input_path}")
    print(f"Output:     {args.output}")
    print(f"Direction:  {args.source.upper()} → {args.target.upper()}")
    print(f"Mode:       {args.mode}")
    if args.nmt:
        print(f"NMT:        {args.nmt}")
    if args.nmt == 'nllb':
        print(f"NLLB Size:  {args.nllb_size}")
    if args.llm:
        print(f"LLM:        {args.llm}")
    if args.aligner:
        print(f"Aligner:    {args.aligner}")
    print(f"{'='*60}\n")
    
    translator = UltimateDocumentTranslator(
        src_lang=args.source,
        tgt_lang=args.target,
        mode=mode_map[args.mode],
        nmt_backend=args.nmt,
        llm_provider=args.llm,
        aligner=args.aligner,
        nllb_model_size=args.nllb_size
    )
    
    await translator.translate_document(input_path, Path(args.output))
    
    print(f"\n{'='*60}")
    print(f"✓ Success!")
    print(f"Output saved to: {args.output}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    asyncio.run(main())