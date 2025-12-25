#!/usr/bin/env python3
import argparse
import asyncio
import logging
import os
import re
import sys
import gc
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Any
from enum import Enum
import requests
from tqdm import tqdm

# --- Core Library Diagnostics ---
print("🔍 System Check...")

def check_library(name, import_stmt):
    try:
        exec(import_stmt, globals())
        print(f"  ✓ {name}")
        return True
    except ImportError:
        print(f"  ⊘ {name} (optional)")
        return False
    except Exception as e:
        print(f"  ✗ {name} error: {e}")
        return False

HAS_DOCX = check_library("python-docx", "from docx import Document; from docx.shared import Pt, RGBColor; from docx.text.paragraph import Paragraph; from docx.oxml.shared import OxmlElement; from docx.oxml.ns import qn")
HAS_TORCH = check_library("torch", "import torch")
HAS_CT2 = check_library("CTranslate2", "import ctranslate2; from huggingface_hub import snapshot_download")
HAS_TRANSFORMERS = check_library("Transformers", "from transformers import AutoTokenizer, AutoModelForSeq2SeqLM")
HAS_OPENAI = check_library("OpenAI", "from openai import AsyncOpenAI")
HAS_ANTHROPIC = check_library("Anthropic", "from anthropic import AsyncAnthropic")

# --- Device & Backend Configuration ---
def get_torch_device():
    if not HAS_TORCH: return "cpu"
    import torch
    if torch.cuda.is_available(): return torch.device("cuda")
    if hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
        return torch.device("mps")
    return torch.device("cpu")

def get_ct2_settings():
    """Optimized for M1 Mac (Accelerate CPU) vs NVIDIA (CUDA)."""
    if HAS_TORCH:
        import torch
        if torch.cuda.is_available(): return "cuda", "float16"
    # Mac M1/M2/M3: MUST use 'cpu' for CT2 and 'int8' for peak ARM64 optimization
    return "cpu", "int8"

def check_fast_align():
    script_dir = Path(__file__).parent
    binary_locations = ["../fast_align/build/fast_align", "./fast_align/build/fast_align", "fast_align"]
    for loc in binary_locations:
        path = script_dir / loc if not loc.startswith('/') else Path(loc)
        if os.path.isfile(path) and os.access(path, os.X_OK): return True
    return False

HAS_FAST_ALIGN = check_fast_align()
print(f"  {'✓' if HAS_FAST_ALIGN else '⊘'} fast_align")
print("-" * 60)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============================================================================
# ENUMS FOR CONFIGURATION
# ============================================================================

class TranslationBackend(Enum):
    """Available translation backends"""
    CT2 = "ct2"
    NLLB = "nllb"
    OPENAI = "openai"
    ANTHROPIC = "anthropic"
    OLLAMA = "ollama"
    AUTO = "auto"


class AlignerBackend(Enum):
    """Available alignment backends"""
    LINDAT = "lindat"
    FAST_ALIGN = "fast_align"
    SIMALIGN = "simalign"
    HEURISTIC = "heuristic"
    AUTO = "auto"


class TranslationMode(Enum):
    """Translation modes"""
    NMT_ONLY = "nmt"
    LLM_WITH_ALIGN = "llm-align"
    LLM_WITHOUT_ALIGN = "llm-plain"
    HYBRID = "hybrid"


# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class FormatRun:
    """A run of text with formatting"""
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    font_color: Optional[Tuple[int, int, int]] = None


@dataclass
class TranslatableParagraph:
    """Paragraph with formatting and metadata"""
    runs: List[FormatRun] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def get_text(self) -> str:
        return ''.join(run.text for run in self.runs)
    
    def get_words(self) -> List[str]:
        """Clean tokenization for the ALIGNER only: strips punctuation."""
        import re
        text = self.get_text()
        # Extracts only alphanumeric sequences (e.g., "Theology" from "Theology:")
        return re.findall(r"\w+", text)
    
    def get_formatted_word_indices(self) -> Dict[str, Set[int]]:
        """Maps formatting to clean alphanumeric word indices."""
        formatted = {'italic': set(), 'bold': set(), 'italic_bold': set()}
        text = self.get_text()
        words = self.get_words() # Uses the same list the aligner sees
        
        if not words:
            return formatted
        
        char_to_word = {}
        last_found = 0
        for word_idx, word in enumerate(words):
            # Find the word in the text, starting search after the previous word
            start = text.find(word, last_found)
            if start != -1:
                for i in range(start, start + len(word)):
                    char_to_word[i] = word_idx
                last_found = start + len(word)
        
        char_pos = 0
        for run in self.runs:
            if not run.text:
                continue
            for char in run.text:
                if char_pos in char_to_word:
                    word_idx = char_to_word[char_pos]
                    if not char.isspace():
                        if run.bold and run.italic:
                            formatted['italic_bold'].add(word_idx)
                        elif run.italic:
                            formatted['italic'].add(word_idx)
                        elif run.bold:
                            formatted['bold'].add(word_idx)
                char_pos += 1
        return formatted


# ============================================================================
# TRANSLATION BACKENDS (keeping previous implementations)
# ============================================================================

class CTranslate2Translator:
    """WMT21 Backend: High-quality dense model."""
    
    MODELS = {
        'en_to_x': 'cstr/wmt21ct2_int8',
        'x_to_en': 'cstr/wmt21-ct2-x-en-int8',
    }
    
    SUPPORTED_LANGS = ['de', 'es', 'fr', 'it', 'ja', 'zh', 'ru', 'pt', 'nl', 'cs', 'uk']
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang, self.tgt_lang = src_lang, tgt_lang
        self.available = False
        self.ct2_dev, self.ct2_compute = get_ct2_settings()
        
        if not HAS_CT2: return
        
        if src_lang == 'en' and tgt_lang in self.SUPPORTED_LANGS:
            self.direction, self.tokenizer_name = 'en_to_x', "facebook/wmt21-dense-24-wide-en-x"
        elif tgt_lang == 'en' and src_lang in self.SUPPORTED_LANGS:
            self.direction, self.tokenizer_name = 'x_to_en', "facebook/wmt21-dense-24-wide-x-en"
        else: return
        
        try:
            logger.info(f"Loading WMT21-CT2 ({self.direction})...")
            model_path = self._get_or_download_model()
            self.translator = ctranslate2.Translator(
                model_path, device=self.ct2_dev, compute_type=self.ct2_compute
            )
            self.tokenizer = AutoTokenizer.from_pretrained(self.tokenizer_name)
            self.available = True
            logger.info(f"✓ WMT21 CT2 ready on {self.ct2_dev}")
        except Exception as e:
            logger.error(f"WMT21 initialization failure: {e}")
    
    def _get_or_download_model(self) -> str:
        cache_base = Path.home() / '.cache' / 'huggingface' / 'hub'
        model_repo = self.MODELS[self.direction]
        model_dir = cache_base / f"models--{model_repo.replace('/', '--')}"
        ref_path = model_dir / 'refs' / 'main'
        
        if ref_path.exists():
            with open(ref_path) as f:
                commit_hash = f.read().strip()
            snapshot_path = model_dir / 'snapshots' / commit_hash
            if (snapshot_path / 'model.bin').exists():
                return str(snapshot_path)
        
        logger.info("Downloading CT2 model...")
        return snapshot_download(repo_id=model_repo)
    
    def translate_batch(self, texts: List[str]) -> List[str]:
        if not self.available or not texts:
            return texts
        
        try:
            # 1. Ensure language tokens are correctly formatted
            # WMT21 models expect the target lang as a prefix
            target_prefix = [[self.tokenizer.lang_code_to_token[self.tgt_lang]]] * len(texts)
            
            # 2. Tokenize with specific padding/truncation
            source_tokens = [self.tokenizer.convert_ids_to_tokens(self.tokenizer.encode(t)) for t in texts]
            
            # 3. Translate with proper beam search and repetition penalty
            results = self.translator.translate_batch(
                source_tokens,
                target_prefix=target_prefix,
                beam_size=5,
                max_batch_size=16,
                repetition_penalty=1.2,
                # Prevent the model from just 'copying' the source if it gets confused
                disable_unk=True 
            )
            
            translated = []
            for i, res in enumerate(results):
                # Strip the language token from the start of the result
                tokens = res.hypotheses[0]
                if self.tokenizer.lang_code_to_token[self.tgt_lang] in tokens:
                    tokens = tokens[tokens.index(self.tokenizer.lang_code_to_token[self.tgt_lang]) + 1:]
                
                decoded = self.tokenizer.decode(self.tokenizer.convert_tokens_to_ids(tokens), skip_special_tokens=True)
                translated.append(decoded.strip())
                
            return translated
        except Exception as e:
            logger.error(f"CT2 Critical Error: {e}")
            return texts


class NLLBTranslator:
    """NLLB-200 translator using CTranslate2 for 4x speedup and low memory"""
    
    # Map sizes to specific CTranslate2 optimized repositories
    REPOS = {
        "600M": "JustFrederik/nllb-200-distilled-600M-ct2-int8",
        "1.3B": "OpenNMT/nllb-200-distilled-1.3B-ct2-int8",
        "3.3B": "OpenNMT/nllb-200-3.3B-ct2-int8"
    }

    LANG_CODES = {
        'en': 'eng_Latn', 'de': 'deu_Latn', 'fr': 'fra_Latn', 
        'es': 'spa_Latn', 'it': 'ita_Latn', 'pt': 'por_Latn',
        'ru': 'rus_Cyrl', 'zh': 'zho_Hans', 'ja': 'jpn_Jpan',
        'ko': 'kor_Hang', 'ar': 'arb_Arab', 'hi': 'hin_Deva',
        'nl': 'nld_Latn', 'pl': 'pol_Latn', 'tr': 'tur_Latn',
        'cs': 'ces_Latn', 'uk': 'ukr_Cyrl', 'vi': 'vie_Latn',
    }
    
    def __init__(self, src_lang: str, tgt_lang: str, model_size: str = "600M"):
        self.src_lang, self.tgt_lang = src_lang, tgt_lang
        self.available = False
        self.mode = None
        self.device = get_torch_device()
        self.ct2_dev, self.ct2_compute = get_ct2_settings()
        
        self.src_code = self.LANG_CODES.get(src_lang)
        self.tgt_code = self.LANG_CODES.get(tgt_lang)
        
        if not self.src_code or not self.tgt_code: return

        ct2_repo = self.REPOS.get(model_size, self.REPOS["600M"])
        standard_repo = f"facebook/nllb-200-distilled-{model_size}"

        if HAS_CT2:
            try:
                logger.info(f"Loading NLLB-CT2 from {ct2_repo}...")
                model_path = snapshot_download(repo_id=ct2_repo)
                try:
                    self.translator = ctranslate2.Translator(
                        model_path, device=self.ct2_dev, compute_type=self.ct2_compute
                    )
                except Exception:
                    # Specific fallback for Mac CPU optimization mismatch
                    self.translator = ctranslate2.Translator(
                        model_path, device="cpu", compute_type="int8"
                    )

                self.tokenizer = AutoTokenizer.from_pretrained(standard_repo, src_lang=self.src_code)
                self.mode, self.available = "ct2", True
                logger.info(f"✓ NLLB-{model_size} CT2 ready.")
                return
            except Exception as e:
                logger.warning(f"NLLB-CT2 failed: {e}")

        if HAS_TRANSFORMERS and HAS_TORCH:
            try:
                logger.info(f"Fallback: Loading standard PyTorch NLLB...")
                self.model = AutoModelForSeq2SeqLM.from_pretrained(standard_repo)
                self.tokenizer = HFTokenizer.from_pretrained(standard_repo, src_lang=self.src_code)
                if self.device.type == "mps": self.model = self.model.half()
                self.model.to(self.device).eval()
                self.mode, self.available = "torch", True
                logger.info(f"✓ NLLB PyTorch ready on {self.device}")
            except Exception as e:
                logger.error(f"Critical: All NLLB paths failed: {e}")
    
    def translate_batch(self, texts: List[str], batch_size: int = 16) -> List[str]:
        """Translate batch using CT2 efficient beam search"""
        if not self.available or not texts:
            return texts
        
        try:
            results = []
            # NLLB requires the target language code as the first token (target prefix)
            target_prefix = [self.tgt_code]
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                
                # Tokenize
                source_tokens = [
                    self.tokenizer.convert_ids_to_tokens(self.tokenizer.encode(t)) 
                    for t in batch
                ]
                
                # Translate
                step_results = self.translator.translate_batch(
                    source_tokens,
                    target_prefix=[target_prefix] * len(batch),
                    beam_size=4,
                    max_batch_size=batch_size,
                    repetition_penalty=1.1
                )
                
                # Decode (skipping the lang code prefix in the output)
                for res in step_results:
                    tokens = res.hypotheses[0]
                    # The first token is usually the target lang code
                    if tokens[0] == self.tgt_code:
                        tokens = tokens[1:]
                    
                    decoded = self.tokenizer.decode(
                        self.tokenizer.convert_tokens_to_ids(tokens), 
                        skip_special_tokens=True
                    )
                    results.append(decoded.strip())
            
            return results
        except Exception as e:
            logger.error(f"NLLB-CT2 translation failed: {e}")
            return texts
        
class LLMTranslator:
    """LLM translator (OpenAI/Anthropic/Ollama)"""
    
    def __init__(self, src_lang: str, tgt_lang: str, preferred_provider: Optional[str] = None):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.providers = self._init_providers(preferred_provider)
        
        if self.providers:
            logger.info(f"✓ LLM available ({list(self.providers.keys())})")
    
    def _init_providers(self, preferred: Optional[str] = None) -> Dict[str, Any]:
        providers = {}
        
        if HAS_OPENAI and os.getenv("OPENAI_API_KEY"):
            providers["openai"] = {
                "client": AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY")),
                "model": os.getenv("OPENAI_MODEL", "gpt-4o-mini")
            }
        
        if HAS_ANTHROPIC and os.getenv("ANTHROPIC_API_KEY"):
            providers["anthropic"] = {
                "client": AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY")),
                "model": os.getenv("ANTHROPIC_MODEL", "claude-3-5-sonnet-20241022")
            }
        
        if self._check_ollama():
            providers["ollama"] = {
                "url": "http://localhost:11434/api/generate",
                "model": self._get_ollama_model()
            }
        
        if preferred and preferred in providers:
            return {preferred: providers[preferred]}
        
        return providers
    
    def _check_ollama(self) -> bool:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            return r.status_code == 200 and len(r.json().get("models", [])) > 0
        except:
            return False
    
    def _get_ollama_model(self) -> str:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            models = r.json().get("models", [])
            return models[0]["name"] if models else "llama3.2"
        except:
            return "llama3.2"
    
    async def translate_text(self, text: str, use_alignment: bool = True) -> Optional[str]:
        """Translate single text"""
        if not text.strip() or not self.providers:
            return None
        
        if use_alignment:
            prompt = (
                f"Translate the following text from {self.src_lang} to {self.tgt_lang}. "
                f"Preserve the word order as much as possible for alignment purposes. "
                f"Return ONLY the translation:\n\n{text}"
            )
        else:
            prompt = (
                f"Translate the following text from {self.src_lang} to {self.tgt_lang}. "
                f"Provide a natural, fluent translation. "
                f"Return ONLY the translation:\n\n{text}"
            )
        
        for provider_name, provider in self.providers.items():
            try:
                if provider_name == "openai":
                    response = await provider["client"].chat.completions.create(
                        model=provider["model"],
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                        max_tokens=4000
                    )
                    return response.choices[0].message.content.strip()
                
                elif provider_name == "anthropic":
                    response = await provider["client"].messages.create(
                        model=provider["model"],
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                        max_tokens=4000
                    )
                    return response.content[0].text.strip()
                
                elif provider_name == "ollama":
                    r = requests.post(
                        provider["url"],
                        json={"model": provider["model"], "prompt": prompt, "stream": False},
                        timeout=120
                    )
                    if r.status_code == 200:
                        return r.json().get("response", "").strip()
            
            except Exception as e:
                logger.debug(f"{provider_name} failed: {e}")
                continue
        
        return None
    
    async def translate_batch(self, texts: List[str], use_alignment: bool = True) -> List[str]:
        """Translate batch"""
        tasks = [self.translate_text(text, use_alignment) for text in texts]
        results = await asyncio.gather(*tasks)
        return [res if res else text for res, text in zip(results, texts)]


# ============================================================================
# WORD ALIGNERS
# ============================================================================

class LindatAligner:
    """Lindat word alignment API"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = self._check_available()
        if self.available:
            logger.info(f"✓ Lindat aligner available ({src_lang}-{tgt_lang})")
        else:
            logger.info(f"⊘ Lindat aligner not available ({src_lang}-{tgt_lang})")
    
    def _check_available(self) -> bool:
        try:
            r = requests.get(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                timeout=5
            )
            return r.status_code in [200, 405]
        except:
            return False
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """
        Extracts high-precision word alignments using BERT embeddings.
        Compatible with all CTranslate2 versions. Uses Softmax Intersection for precision.
        """
        if not self.available or not src_words or not tgt_words:
            return []
        
        # Local imports for total robustness
        import numpy as np
        try:
            # 1. PRE-PROCESSING
            def get_tokens_and_map(words):
                subtokens, word_map = [], []
                for i, w in enumerate(words):
                    tokens = self.tokenizer.tokenize(w) or [self.tokenizer.unk_token]
                    subtokens.extend(tokens)
                    word_map.extend([i] * len(tokens))
                return subtokens, word_map

            src_subtokens, src_word_map = get_tokens_and_map(src_words)
            tgt_subtokens, tgt_word_map = get_tokens_and_map(tgt_words)

            # 2. EMBEDDING EXTRACTION
            if self.mode == "ct2":
                # Prepare inputs
                src_input = [["[CLS]"] + src_subtokens + ["[SEP]"]]
                tgt_input = [["[CLS]"] + tgt_subtokens + ["[SEP]"]]
                
                # Standard forward (compatible with all CT2 versions)
                res_src = self.encoder.forward_batch(src_input)
                res_tgt = self.encoder.forward_batch(tgt_input)
                
                # Extract last hidden state (Layer 12)
                src_out = np.array(res_src.last_hidden_state)[0, 1:-1]
                tgt_out = np.array(res_tgt.last_hidden_state)[0, 1:-1]
            else:
                import torch
                def to_ids(tokens):
                    ids = self.tokenizer.convert_tokens_to_ids(tokens)
                    return torch.tensor([self.tokenizer.cls_token_id] + ids + [self.tokenizer.sep_token_id]).to(self.device)
                
                with torch.no_grad():
                    # PyTorch still targets Layer 8 for maximum precision
                    src_h = self.model(to_ids(src_subtokens).unsqueeze(0), output_hidden_states=True)[2][8][0, 1:-1]
                    tgt_h = self.model(to_ids(tgt_subtokens).unsqueeze(0), output_hidden_states=True)[2][8][0, 1:-1]
                    src_out = src_h.detach().cpu().float().numpy()
                    tgt_out = tgt_h.detach().cpu().float().numpy()

            # 3. SIMILARITY & SOFTMAX INTERSECTION
            # Normalize
            src_norm = src_out / np.linalg.norm(src_out, axis=-1, keepdims=True)
            tgt_norm = tgt_out / np.linalg.norm(tgt_out, axis=-1, keepdims=True)
            similarity = np.dot(src_norm, tgt_norm.T)

            # Softmax Intersection Logic
            def softmax(x, axis):
                e_x = np.exp(x - np.max(x, axis=axis, keepdims=True))
                return e_x / e_x.sum(axis=axis, keepdims=True)

            probs_src_to_tgt = softmax(similarity, axis=1)
            probs_tgt_to_src = softmax(similarity, axis=0)
            
            # Mutual agreement threshold
            threshold = 1e-3
            mask = (probs_src_to_tgt > threshold) * (probs_tgt_to_src > threshold)
            
            # 4. MAPPING
            coords = np.argwhere(mask)
            align_words = set()
            for i, j in coords:
                align_words.add((src_word_map[i], tgt_word_map[j]))
            
            final_alignments = sorted(list(align_words))
            
            # VERBOSE CLI LOG
            logger.debug(f"TRACE | Awesome-Align ({self.mode.upper()}) | Clean Words: {len(src_words)}x{len(tgt_words)} | Links Found: {len(final_alignments)}")

            return final_alignments
            
        except Exception as e:
            logger.error(f"ALIGN | Critical Logic Error: {e}")
            return []

class AwesomeAlignAligner:
    """BERT Aligner: Uses custom CT2-INT8 model from HuggingFace."""
    
    def __init__(self):
        self.available = False
        self.mode = None
        self.device = get_torch_device()
        self.ct2_dev, self.ct2_compute = get_ct2_settings()
        self.ct2_repo = "cstr/bert-base-multilingual-cased-ct2-int8"
        self.standard_repo = "bert-base-multilingual-cased"

        if HAS_CT2:
            try:
                logger.info(f"Loading CT2 Aligner from {self.ct2_repo}...")
                model_path = snapshot_download(repo_id=self.ct2_repo)
                self.encoder = ctranslate2.Encoder(
                    model_path, device=self.ct2_dev, compute_type=self.ct2_compute,
                    intra_threads=1 
                )
                # Load tokenizer from string to avoid local JSON collision
                self.tokenizer = AutoTokenizer.from_pretrained(self.standard_repo)
                self.mode, self.available = "ct2", True
                logger.info("✓ Awesome-Align CT2 ready.")
                return
            except Exception as e:
                logger.warning(f"CT2 Aligner load failed: {e}")

        if HAS_TRANSFORMERS and HAS_TORCH:
            try:
                from transformers import BertModel, BertTokenizer
                self.model = BertModel.from_pretrained(self.standard_repo)
                self.tokenizer = BertTokenizer.from_pretrained(self.standard_repo)
                if self.device.type == "mps": self.model = self.model.half()
                self.model.to(self.device).eval()
                self.mode, self.available = "torch", True
                logger.info(f"✓ Awesome-Align PyTorch ready on {self.device}")
            except Exception as e:
                logger.error(f"Critical: Aligner fallback failed: {e}")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """
        Extracts high-precision word alignments using BERT embeddings.
        Uses Mutual Argmax (Intersection) logic for 1-to-1 precision.
        Compatible with Mac CTranslate2 (no return_all_layers).
        """
        if not self.available or not src_words or not tgt_words:
            return []
        
        import numpy as np
        src_out, tgt_out = None, None
        
        try:
            # 1. PRE-PROCESSING: Subword tokenization
            def get_tokens_and_map(words):
                subtokens, word_map = [], []
                for i, w in enumerate(words):
                    tokens = self.tokenizer.tokenize(w) or [self.tokenizer.unk_token]
                    subtokens.extend(tokens)
                    word_map.extend([i] * len(tokens))
                return subtokens, word_map

            src_subtokens, src_word_map = get_tokens_and_map(src_words)
            tgt_subtokens, tgt_word_map = get_tokens_and_map(tgt_words)

            # 2. EMBEDDING EXTRACTION
            if self.mode == "ct2":
                src_input = [["[CLS]"] + src_subtokens + ["[SEP]"]]
                tgt_input = [["[CLS]"] + tgt_subtokens + ["[SEP]"]]
                
                # Use standard batch forward for cross-version compatibility
                res_src = self.encoder.forward_batch(src_input)
                res_tgt = self.encoder.forward_batch(tgt_input)
                
                # Extract Layer 12 (last_hidden_state)
                src_out = np.array(res_src.last_hidden_state)[0, 1:-1]
                tgt_out = np.array(res_tgt.last_hidden_state)[0, 1:-1]
                
            else: # PyTorch Mode
                import torch
                def to_ids(tokens):
                    ids = self.tokenizer.convert_tokens_to_ids(tokens)
                    return torch.tensor([self.tokenizer.cls_token_id] + ids + [self.tokenizer.sep_token_id]).to(self.device)
                
                with torch.no_grad():
                    # PyTorch uses Layer 8 (sweet spot)
                    out_s = self.model(to_ids(src_subtokens).unsqueeze(0), output_hidden_states=True)[2][8][0, 1:-1]
                    out_t = self.model(to_ids(tgt_subtokens).unsqueeze(0), output_hidden_states=True)[2][8][0, 1:-1]
                    src_out = out_s.detach().cpu().float().numpy()
                    tgt_out = out_t.detach().cpu().float().numpy()

            # 3. STABLE ALIGNMENT LOGIC: Mutual Argmax
            # Normalize vectors for cosine similarity
            src_norm = src_out / np.linalg.norm(src_out, axis=-1, keepdims=True)
            tgt_norm = tgt_out / np.linalg.norm(tgt_out, axis=-1, keepdims=True)
            similarity = np.dot(src_norm, tgt_norm.T)

            # Find best matches in both directions
            best_tgt_for_src = np.argmax(similarity, axis=1) # Shape: (src_len,)
            best_src_for_tgt = np.argmax(similarity, axis=0) # Shape: (tgt_len,)

            threshold = 1e-3
            align_words = set()

            # Mutual Agreement (The standard working method)
            for i, j in enumerate(best_tgt_for_src):
                # If source i picked target j, AND target j picked source i...
                if best_src_for_tgt[j] == i and similarity[i, j] > threshold:
                    # Map subword indices back to word indices
                    align_words.add((src_word_map[i], tgt_word_map[j]))
            
            final_alignments = sorted(list(align_words))
            
            # VERBOSE CLI LOG
            logger.debug(f"TRACE | Awesome-Align | Mode: {self.mode.upper()} | Links: {len(final_alignments)}")

            return final_alignments
            
        except Exception as e:
            logger.error(f"ALIGN | Logic failure: {e}")
            return []
    
    def align_old(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            import numpy as np
            
            # 1. Tokenize
            def get_tokens_and_map(words):
                tokens = []
                word_map = []
                for i, w in enumerate(words):
                    subwords = self.tokenizer.tokenize(w)
                    tokens.extend(subwords)
                    word_map.extend([i] * len(subwords))
                return tokens, word_map

            src_tokens, src_map = get_tokens_and_map(src_words)
            tgt_tokens, tgt_map = get_tokens_and_map(tgt_words)

            # 2. Extract Embeddings using CTranslate2
            # We add [CLS] and [SEP] just like BERT expects
            src_input = [["[CLS]"] + src_tokens + ["[SEP]"]]
            tgt_input = [["[CLS]"] + tgt_tokens + ["[SEP]"]]

            # forward_batch returns a StorageView; we convert to numpy for easy math
            src_out = np.array(self.encoder.forward_batch(src_input).last_hidden_state)[0, 1:-1]
            tgt_out = np.array(self.encoder.forward_batch(tgt_input).last_hidden_state)[0, 1:-1]

            # 3. Compute Similarity (Dot Product)
            # Normalizing vectors first ensures we are doing Cosine Similarity
            src_out /= np.linalg.norm(src_out, axis=-1, keepdims=True)
            tgt_out /= np.linalg.norm(tgt_out, axis=-1, keepdims=True)
            
            similarity = np.dot(src_out, tgt_out.T)

            # 4. Extract Alignments (Mutual Argmax / Threshold)
            threshold = 1e-3
            best_src = np.argmax(similarity, axis=1)
            best_tgt = np.argmax(similarity, axis=0)

            align_words = set()
            for i, j in enumerate(best_src):
                if best_tgt[j] == i and similarity[i, j] > threshold:
                    align_words.add((src_map[i], tgt_map[j]))
            
            return sorted(list(align_words))
            
        except Exception as e:
            logger.debug(f"CT2 Alignment failed: {e}")
            return []

class FastAlignAligner:
    """fast_align local aligner - uses temp file for binary"""
    
    def __init__(self):
        self.available = False
        self.mode = None
        self.binary_path = None
        self.atools_path = None
        
        # Check for Python package first
        try:
            import fast_align
            self.available = True
            self.mode = "python"
            return
        except ImportError:
            pass
        
        # Binary search logic
        script_dir = Path(__file__).parent
        search_dirs = [script_dir / "../fast_align/build", script_dir / "./fast_align/build"]
        
        for d in search_dirs:
            fa = d / "fast_align"
            at = d / "atools"
            if fa.exists() and os.access(fa, os.X_OK):
                self.binary_path = str(fa)
                self.atools_path = str(at) if at.exists() else None
                self.available = True
                self.mode = "binary"
                return
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            if self.mode == "python":
                from fast_align import align
                src_text = ' '.join(src_words)
                tgt_text = ' '.join(tgt_words)
                result = align([f"{src_text} ||| {tgt_text}"], forward=True)
                return [tuple(map(int, p.split('-'))) for p in result[0].split()]
            
            elif self.mode == "binary":
                import subprocess
                import tempfile
                
                input_str = f"{' '.join(src_words)} ||| {' '.join(tgt_words)}\n"
                
                with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
                    f.write(input_str)
                    temp_path = f.name
                
                try:
                    # -d: diagonal, -o: optimize tension, -v: variational bayes
                    # -I 1: single pass is enough for alignment of a known translation
                    cmd = [self.binary_path, '-i', temp_path, '-d', '-o', '-v', '-I', '1']
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                    
                    if result.returncode != 0:
                        return []
                    
                    # fast_align binary outputs to stdout in Pharaoh format (i-j)
                    alignments = []
                    output = result.stdout.strip()
                    if output:
                        for pair in output.split():
                            if '-' in pair:
                                i, j = map(int, pair.split('-'))
                                alignments.append((i, j))
                    return alignments
                finally:
                    if os.path.exists(temp_path): os.unlink(temp_path)
        except Exception as e:
            logger.debug(f"fast_align execution failed: {e}")
            return []


class SimAlignAligner:
    """SimAlign aligner"""
    
    def __init__(self):
        self.available = HAS_SIMALIGN
        if self.available:
            try:
                self.aligner = SentenceAligner(model="bert", token_type="bpe", matching_methods="mai")
                logger.info("✓ SimAlign available")
            except Exception as e:
                logger.debug(f"SimAlign init failed: {e}")
                self.available = False
                logger.info("⊘ SimAlign init failed")
        else:
            logger.info("⊘ SimAlign not available (install: pip install simalign)")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            result = self.aligner.get_word_aligns(src_words, tgt_words)
            alignments = []
            for src_idx, tgt_idx in result['mwmf'].items():
                alignments.append((src_idx, tgt_idx))
            return alignments
        except Exception as e:
            logger.debug(f"SimAlign failed: {e}")
            return []


class HeuristicAligner:
    """Heuristic fallback - align words that appear in both"""
    
    def __init__(self):
        logger.info("✓ Heuristic aligner (fallback)")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Simple heuristic alignment"""
        alignments = []
        
        src_lower = [w.lower().strip('.,!?;:') for w in src_words]
        tgt_lower = [w.lower().strip('.,!?;:') for w in tgt_words]
        
        for i, src_word in enumerate(src_lower):
            for j, tgt_word in enumerate(tgt_lower):
                if src_word == tgt_word and len(src_word) > 2:
                    alignments.append((i, j))
        
        return alignments


class MultiAligner:
    """Try multiple aligners with proper fallback"""
    
    def __init__(self, src_lang: str, tgt_lang: str, preferred: Optional[str] = None):
        self.aligners = []
        
        # If a specific aligner was requested
        if preferred and preferred != "auto":
            if preferred == "lindat":
                lindat = LindatAligner(src_lang, tgt_lang)
                if lindat.available:
                    self.aligners.append(("Lindat", lindat))
                else:
                    logger.warning(f"Requested aligner '{preferred}' not available, will try others")
            
            elif preferred == "awesome":
                awesome = AwesomeAlignAligner()
                if awesome.available:
                    self.aligners.append(("awesome-align", awesome))
            
            elif preferred == "fast_align":
                fast_align = FastAlignAligner()
                if HAS_FAST_ALIGN: 
                    self.aligners.append(("fast_align", fast_align))
                else:
                    logger.warning(f"fast_align requested but binary/package not found.")

            elif preferred == "simalign":
                simalign = SimAlignAligner()
                if simalign.available:
                    self.aligners.append(("SimAlign", simalign))
                else:
                    logger.warning(f"Requested aligner '{preferred}' not available, will try others")
            
            elif preferred == "heuristic":
                self.aligners.append(("Heuristic", HeuristicAligner()))
        
        # Auto mode or fallback: try all available aligners
        if not self.aligners or preferred == "auto":
            # Try Lindat first
            lindat = LindatAligner(src_lang, tgt_lang)
            if lindat.available:
                self.aligners.append(("Lindat", lindat))

            awesome = AwesomeAlignAligner()
            if awesome.available:
                self.aligners.append(("awesome-align", awesome))
            
            # Then SimAlign
            simalign = SimAlignAligner()
            if simalign.available:
                self.aligners.append(("SimAlign", simalign))
            
            # Then fast_align but will work right only with snapshot
            fast_align = FastAlignAligner()
            if fast_align.available:
                self.aligners.append(("fast_align", fast_align))
        
        # Always add heuristic as final fallback
        if not any(name == "Heuristic" for name, _ in self.aligners):
            self.aligners.append(("Heuristic", HeuristicAligner()))
        
        logger.info(f"Aligner chain: {[name for name, _ in self.aligners]}")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Try aligners in order until one succeeds"""
        for name, aligner in self.aligners:
            result = aligner.align(src_words, tgt_words)
            if result:
                logger.debug(f"✓ {name} alignment: {len(result)} links")
                return result
        
        logger.debug("Using heuristic fallback (no quality alignments found)")
        return []


# ============================================================================
# DOCUMENT TRANSLATOR
# ============================================================================

class UltimateDocumentTranslator:
    """Document translator with configurable backends and M1 optimization."""
    
    def __init__(
        self,
        src_lang: str,
        tgt_lang: str,
        mode: TranslationMode = TranslationMode.HYBRID,
        nmt_backend: Optional[str] = "nllb",
        llm_provider: Optional[str] = None,
        aligner: Optional[str] = None,
        nllb_model_size: str = "600M"
    ):
        self.src_lang, self.tgt_lang, self.mode = src_lang, tgt_lang, mode
        self.ct2 = None  # WMT translator
        self.nllb = None # NLLB translator
        self.llm = None
        self.aligner = None

        logger.info(f"INIT | Starting Translator ({src_lang}→{tgt_lang})")
        logger.info(f"INIT | Mode: {mode.value} | NMT: {nmt_backend} | Aligner: {aligner or 'auto'}")
        self.log_memory("Initialization Start")

        # 1. NMT BACKEND SELECTION (Exclusive)
        if mode in [TranslationMode.NMT_ONLY, TranslationMode.HYBRID]:
            if nmt_backend == "nllb":
                self.nllb = NLLBTranslator(src_lang, tgt_lang, nllb_model_size)
                if not self.nllb.available:
                    logger.error("INIT | CRITICAL: NLLB specifically requested but failed to load.")
            
            elif nmt_backend == "ct2":
                self.ct2 = CTranslate2Translator(src_lang, tgt_lang)
                if not self.ct2.available:
                    logger.error("INIT | CRITICAL: CTranslate2/WMT specifically requested but failed to load.")
            
            elif nmt_backend == "auto" or nmt_backend is None:
                # Priority: NLLB (Lightweight) -> CT2 (Dense)
                self.nllb = NLLBTranslator(src_lang, tgt_lang, nllb_model_size)
                if not (self.nllb and self.nllb.available):
                    logger.info("INIT | NLLB unavailable, trying CTranslate2/WMT...")
                    self.ct2 = CTranslate2Translator(src_lang, tgt_lang)

        # 2. LLM BACKEND (Hybrid or LLM modes)
        if mode in [TranslationMode.LLM_WITH_ALIGN, TranslationMode.LLM_WITHOUT_ALIGN, TranslationMode.HYBRID]:
            self.llm = LLMTranslator(src_lang, tgt_lang, llm_provider)
            if not self.llm.providers:
                logger.warning("INIT | Mode requires LLM but no providers (OpenAI/Anthropic/Ollama) available.")

        # 3. ALIGNER SELECTION (Priority: Awesome-Align)
        if mode in [TranslationMode.LLM_WITH_ALIGN, TranslationMode.HYBRID, TranslationMode.NMT_ONLY]:
            # We force 'awesome' as the preferred auto-aligner for M1 precision
            aligner_choice = aligner if aligner and aligner != "auto" else "awesome"
            self.aligner = MultiAligner(src_lang, tgt_lang, aligner_choice)
            
        self.log_memory("Initialization Complete")

    def log_memory(self, stage: str):
        """Log current RAM usage (requires psutil)."""
        try:
            import psutil
            process = psutil.Process(os.getpid())
            mem_mb = process.memory_info().rss / (1024 * 1024)
            logger.debug(f"MEM | Stage: {stage} | Usage: {mem_mb:.2f} MB")
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"Memory log failed: {e}")
    
    async def translate_text(self, text: str) -> str:
        if not text.strip(): return text
        
        # 1. Try NLLB if loaded
        if self.nllb and self.nllb.available:
            result = self.nllb.translate_batch([text])[0]
            if result and result != text: return result
            
        # 2. Try WMT21 if loaded
        if self.ct2 and self.ct2.available:
            result = self.ct2.translate_batch([text])[0]
            if result and result != text: return result

        # 3. Try LLM (Hybrid/LLM modes)
        if self.llm and self.llm.providers:
            # use_alignment depends on mode
            use_align = (self.mode != TranslationMode.LLM_WITHOUT_ALIGN)
            result = await self.llm.translate_text(text, use_alignment=use_align)
            if result: return result
            
        return text
    
    def extract_paragraph(self, para: Paragraph) -> TranslatableParagraph:
        """Extracts paragraph with resolved font hierarchy to prevent Theme reversion."""
        # 1. Resolve the "Base Font" for this paragraph
        def get_resolved_font_name(p):
            # Check runs first
            for r in p.runs:
                if r.font.name: return r.font.name
            # Check style hierarchy
            curr_style = p.style
            while curr_style:
                if curr_style.font.name: return curr_style.font.name
                curr_style = curr_style.base_style
            return "Times New Roman" # Fallback if everything is 'None'

        resolved_base_font = get_resolved_font_name(para)
        
        runs = []
        for run in para.runs:
            f_color = None
            try:
                if run.font.color and run.font.color.rgb:
                    f_color = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
            except: pass

            runs.append(FormatRun(
                text=run.text,
                bold=run.bold,
                italic=run.italic,
                underline=run.underline,
                # If run has no font name, use the resolved base font we found
                font_name=run.font.name if run.font.name else resolved_base_font,
                font_size=run.font.size.pt if run.font.size else (para.style.font.size.pt if para.style.font.size else 12.0),
                font_color=f_color
            ))
        
        trans_para = TranslatableParagraph(runs=runs)
        trans_para.metadata['style'] = para.style
        trans_para.metadata['alignment'] = para.alignment
        
        pf = para.paragraph_format
        trans_para.metadata['layout'] = {
            'left_indent': pf.left_indent,
            'right_indent': pf.right_indent,
            'first_line_indent': pf.first_line_indent,
            'line_spacing': pf.line_spacing,
            'space_before': pf.space_before,
            'space_after': pf.space_after
        }
        return trans_para
    
    def log_memory(self, stage: str):
        """Log current RAM usage (requires psutil)."""
        try:
            import psutil
            process = psutil.Process(os.getpid())
            mem_mb = process.memory_info().rss / (1024 * 1024)
            logger.debug(f"MEM | Stage: {stage} | Usage: {mem_mb:.2f} MB")
        except ImportError:
            pass

    def copy_font_properties(self, target_run, source_run: FormatRun):
        """Forces Font Name, Size, and Color into XML. Does NOT touch bold/italic."""
        try:
            if source_run.font_size:
                target_run.font.size = Pt(source_run.font_size)
            if source_run.font_color:
                target_run.font.color.rgb = RGBColor(*source_run.font_color)
            if source_run.underline is not None:
                target_run.font.underline = source_run.underline

            if source_run.font_name:
                # The XML 'rFonts' injection to bypass Theme defaults
                r = target_run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:ascii'), source_run.font_name)
                rFonts.set(qn('w:hAnsi'), source_run.font_name)
                rFonts.set(qn('w:eastAsia'), source_run.font_name)
                rFonts.set(qn('w:cs'), source_run.font_name)
                target_run.font.name = source_run.font_name
        except Exception as e:
            logger.debug(f"TRACE | Font Force Error: {e}")
    
    def apply_aligned_formatting(self, para: Paragraph, trans_para: TranslatableParagraph, translated_text: str, alignment: List[Tuple[int, int]]):
        """Reconstructs paragraph using a precise map of inline styles and anchored footnotes."""
        self.log_para_trace(para, "INPUT")
        p_element = para._p
        footnote_refs = p_element.xpath('.//w:r[w:footnoteReference or w:footnoteRef]')
        
        # 1. Restore Para-level metadata
        if trans_para.metadata.get('style'):
            para.style = trans_para.metadata['style']
        para.alignment = trans_para.metadata.get('alignment')
        
        # Re-apply layout spacing/indents
        layout = trans_para.metadata.get('layout', {})
        pf = para.paragraph_format
        for attr, val in layout.items():
            try:
                if val is not None: setattr(pf, attr, val)
            except: pass

        # 2. Clear content
        for run in para.runs:
            p_element.remove(run._element)

        # 3. Footnote text paragraphs: re-attach numbering at start
        is_footnote_para = "footnote" in str(para.style.name).lower()
        if is_footnote_para and footnote_refs:
            for ref in footnote_refs: p_element.append(ref)
            para.add_run("\u00A0") 

        # 4. Map formatting to Clean Target Indices
        src_clean_words = trans_para.get_words()
        tgt_raw_units = translated_text.split() 
        formatted_indices = trans_para.get_formatted_word_indices()
        
        clean_to_raw_tgt = {}
        clean_idx = 0
        for raw_idx, unit in enumerate(tgt_raw_units):
            if re.search(r'\w', unit):
                clean_to_raw_tgt[clean_idx] = raw_idx
                clean_idx += 1

        # Use the first run as the baseline "Aesthetic" template (Font, Size, Color)
        font_template = trans_para.runs[0] if trans_para.runs else None

        # 5. Build units with Inline Style merging
        for i, unit in enumerate(tgt_raw_units):
            run_text = unit + (" " if i < len(tgt_raw_units)-1 else "")
            run = para.add_run(run_text)
            
            # Determine Bold/Italic/Underline for this unit via Aligner
            style_type = None
            for src_idx, align_tgt_idx in alignment:
                if clean_to_raw_tgt.get(align_tgt_idx) == i:
                    for s in ['italic_bold', 'bold', 'italic']:
                        if src_idx in formatted_indices[s]:
                            style_type = s; break
            
            # Apply Aligned Inline Styles
            if style_type == 'italic_bold': run.bold = run.italic = True
            elif style_type == 'bold': run.bold = True
            elif style_type == 'italic': run.italic = True
            
            # Apply Baseline aesthetics (Font name, size, color)
            if font_template:
                self.copy_font_properties(run, font_template)

        # 6. Re-attach body footnotes to the end of the translated sentence
        if not is_footnote_para and footnote_refs:
            logger.debug(f"TRACE | Body Footnote | Re-anchoring {len(footnote_refs)} refs.")
            for ref in footnote_refs: p_element.append(ref)

        self.log_para_trace(para, "OUTPUT")
    
    def is_paragraph_safe_to_translate(self, para: Paragraph) -> bool:
        """Check if paragraph can be safely translated"""
        if not para.text or not para.text.strip():
            return False
        
        text = para.text.strip()
        if not text:
            return False
        
        if len(para.text.strip()) <= 1 and not any(run.text.strip() for run in para.runs):
            return False
        
        try:
            for run in para.runs:
                if run._element.xpath('.//w:drawing | .//w:pict'):
                    logger.debug(f"Skipping paragraph with drawing/image")
                    return False
        except:
            pass
        
        try:
            field_chars = para._element.xpath('.//w:fldChar')
            if field_chars:
                is_footnote_field = para._element.xpath('.//w:footnoteReference')
                if not is_footnote_field:
                    logger.debug(f"Skipping paragraph with field")
                    return False
        except:
            pass
        
        return True
    
    async def translate_paragraph(self, para: Paragraph):
        """Paragraph translation with heavy trace logging for format debugging."""
        if not para.text or not para.text.strip() or not self.is_paragraph_safe_to_translate(para):
            return
        
        try:
            self.log_memory("Pre-Paragraph")
            trans_para = self.extract_paragraph(para)
            original_text = trans_para.get_text()
            
            sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', original_text) if s.strip()]
            if not sentences: return

            if self.ct2 and self.ct2.available:
                translated_sentences = self.ct2.translate_batch(sentences)
            elif self.nllb and self.nllb.available:
                translated_sentences = self.nllb.translate_batch(sentences)
            else: return

            translated_text = " ".join(translated_sentences)
            
            # Prep for Aligner
            src_words = trans_para.get_words() 
            tgt_words_clean = re.findall(r"\w+", translated_text)
            
            logger.debug("-" * 40)
            logger.debug(f"TRACE | RAW SRC: {original_text[:100]}...")
            logger.debug(f"TRACE | CLEAN SRC WORDS: {src_words}")
            logger.debug(f"TRACE | CLEAN TGT WORDS: {tgt_words_clean}")
            
            if not src_words or not tgt_words_clean:
                para.text = translated_text
                return
                
            alignment = self.aligner.align(src_words, tgt_words_clean)
            logger.debug(f"TRACE | ALIGNMENT MAP: {alignment}")
            
            self.apply_aligned_formatting(para, trans_para, translated_text, alignment)
            logger.debug("-" * 40)
            self.log_memory("Post-Paragraph")
                    
        except Exception as e:
            logger.error(f"Error in translate_paragraph: {e}", exc_info=True)
    
    def get_footnotes(self, doc: Document) -> List[Paragraph]:
        """Extract footnotes. We pass doc as parent to avoid Part attribute error"""
        try:
            document_part = doc.part
            footnote_part = None
            
            for rel in document_part.rels.values():
                if "relationships/footnotes" in rel.reltype:
                    footnote_part = rel.target_part
                    break
            
            if not footnote_part:
                return []
            
            from docx.oxml import parse_xml
            root = parse_xml(footnote_part.blob)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            paragraphs = []
            for footnote in root.xpath('//w:footnote', namespaces=ns):
                f_id = footnote.get(f'{{{ns["w"]}}}id')
                # Skip internal Word markers (id 0 and -1)
                if f_id and int(f_id) <= 0:
                    continue
                
                for p_elem in footnote.xpath('.//w:p', namespaces=ns):
                    # FIX: Pass 'doc' instead of 'None' so para.part is valid
                    para = Paragraph(p_elem, doc)
                    if para.text.strip():
                        paragraphs.append(para)
            
            self._footnote_root = root
            self._footnote_part = footnote_part
            return paragraphs
            
        except Exception as e:
            logger.warning(f"Could not extract footnotes: {e}")
            return []
    
    def get_all_paragraphs(self, doc: Document) -> List[Tuple[Paragraph, str]]:
        """Aggregate all paragraphs from document"""
        all_paras = []
        
        # Main body
        for para in doc.paragraphs:
            all_paras.append((para, "body"))
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_paras.append((para, "table"))
        
        # Footnotes - with error handling
        try:
            footnote_paras = self.get_footnotes(doc)
            for para in footnote_paras:
                all_paras.append((para, "footnote"))
        except Exception as e:
            logger.warning(f"Skipping footnotes due to error: {e}")
        
        # Headers/Footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                all_paras.append((para, "header"))
            for para in section.footer.paragraphs:
                all_paras.append((para, "footer"))
        
        return all_paras
    
    def log_document_info(self, doc: Document, label: str):
        """Logs global document properties and style inventory."""
        logger.debug(f"{'='*20} DOCUMENT INFO [{label}] {'='*20}")
        for i, section in enumerate(doc.sections):
            logger.debug(f"Section {i} | Size: {section.page_width.pt:.1f}x{section.page_height.pt:.1f}pt")
            logger.debug(f"Section {i} | Margins: L:{section.left_margin.pt:.1f} R:{section.right_margin.pt:.1f} T:{section.top_margin.pt:.1f} B:{section.bottom_margin.pt:.1f}")
            logger.debug(f"Section {i} | Gutter: {section.gutter.pt:.1f}pt | Header-Dist: {section.header_distance.pt:.1f}pt")
        
        # Log all paragraph styles present in the document
        style_names = [s.name for s in doc.styles if s.type == 1]
        logger.debug(f"Style Inventory ({len(style_names)}): {', '.join(style_names)}")

    def log_para_trace(self, para: Paragraph, label: str):
        """Detailed debug log showing style, font, and layout metrics."""
        pf = para.paragraph_format
        style = para.style
        
        # Check if first run has inline formatting
        has_bold = any(r.bold for r in para.runs)
        has_ital = any(r.italic for r in para.runs)
        
        f_name = para.runs[0].font.name if para.runs and para.runs[0].font.name else "Inherited"
        f_size = f"{para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else 'Default'}pt"
        
        logger.debug(f"PARA [{label}] | Style: '{style.name}'")
        logger.debug(f"  > Font: {f_name} @ {f_size} | Bold-Any: {has_bold} | Ital-Any: {has_ital}")
        logger.debug(f"  > Indent-L: {pf.left_indent.pt if pf.left_indent else 0:.1f}pt | Spacing-A: {pf.space_after.pt if pf.space_after else 0:.1f}pt")

    async def translate_document(self, input_path: Path, output_path: Path):
        """Full document lifecycle with robust XML commitment and verification logs."""
        self.log_memory("Initialization")
        doc = Document(str(input_path))
        
        if logger.isEnabledFor(logging.DEBUG):
            self.log_document_info(doc, "INPUT")

        # Gather footnotes properly to initialize _footnote_root
        _ = self.get_footnotes(doc)
        all_paras = self.get_all_paragraphs(doc)
        translatable = [(p, l) for p, l in all_paras if p.text.strip() and self.is_paragraph_safe_to_translate(p)]
        
        logger.info(f"Processing {len(translatable)} paragraphs across {len(all_paras)} total entities.")
        
        for para, location in tqdm(translatable, desc="Translating"):
            try:
                await self.translate_paragraph(para)
            except Exception as e:
                logger.error(f"Error in {location} translation: {e}")

        # FOOTNOTE XML COMMITMENT
        if hasattr(self, '_footnote_part') and hasattr(self, '_footnote_root'):
            try:
                from lxml import etree
                # Standard etree tostring ensures namespace 'w' is preserved correctly
                updated_xml = etree.tostring(self._footnote_root, encoding='utf-8', xml_declaration=True)
                self._footnote_part._blob = updated_xml
                logger.info("✓ Success | Footnote XML blob successfully serialized.")
            except Exception as e:
                logger.error(f"Error | Footnote commitment failed: {e}")

        logger.info(f"Saving file to {output_path}")
        doc.save(str(output_path))
        
        if logger.isEnabledFor(logging.DEBUG):
            self.log_document_info(Document(str(output_path)), "OUTPUT")
        logger.info("✓ Document Translation Complete.")


# ============================================================================
# CLI
# ============================================================================

async def main():
    parser = argparse.ArgumentParser(
        description='Document Translator',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Default (auto-detect best backends)
  %(prog)s input.docx output.docx -s en -t de
  
  # Use NLLB for translation
  %(prog)s input.docx output.docx -s en -t fr --nmt nllb
  
  # Use LLM with alignment
  %(prog)s input.docx output.docx -s en -t es --mode llm-align --llm openai
  
  # Use LLM without alignment (natural translation)
  %(prog)s input.docx output.docx -s en -t it --mode llm-plain --llm anthropic
  
  # Hybrid mode (will use Lindat aligner)
  %(prog)s input.docx output.docx -s en -t de --mode hybrid
  
  # Use larger NLLB model
  %(prog)s input.docx output.docx -s en -t ja --nmt nllb --nllb-size 1.3B

Environment Variables:
  OPENAI_API_KEY       - OpenAI API key
  OPENAI_MODEL         - OpenAI model (default: gpt-4o-mini)
  ANTHROPIC_API_KEY    - Anthropic API key
  ANTHROPIC_MODEL      - Anthropic model (default: claude-3-5-sonnet-20241022)

Installation Notes:
  fast_align: This is a C++ tool, not a Python package. 
              See: https://github.com/clab/fast_align
              Or install Python wrapper: pip install fast-align
  
  SimAlign:   pip install simalign
  NLLB:       pip install transformers torch
        """
    )
    
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('output', help='Output .docx file')
    parser.add_argument('-s', '--source', default='en', help='Source language (default: en)')
    parser.add_argument('-t', '--target', default='de', help='Target language (default: de)')
    
    parser.add_argument(
        '--mode',
        choices=['nmt', 'llm-align', 'llm-plain', 'hybrid'],
        default='hybrid',
        help='Translation mode (default: hybrid)'
    )
    
    parser.add_argument(
        '--nmt',
        choices=['ct2', 'nllb', 'auto'],
        default='nllb',
        help='NMT backend (default: nllb)'
    )
    
    parser.add_argument(
        '--nllb-size',
        choices=['600M', '1.3B', '3.3B'],
        default='600M',
        help='NLLB model size (default: 600M)'
    )
    
    parser.add_argument(
        '--llm',
        choices=['openai', 'anthropic', 'ollama'],
        help='LLM provider (default: auto-detect)'
    )
    
    parser.add_argument(
        '--aligner',
        choices=['lindat', 'fast_align', 'simalign', 'awesome', 'heuristic', 'auto'],
        help='Word aligner (default: auto)'
    )
    
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"Error: Only .docx files supported")
        sys.exit(1)
    
    mode_map = {
        'nmt': TranslationMode.NMT_ONLY,
        'llm-align': TranslationMode.LLM_WITH_ALIGN,
        'llm-plain': TranslationMode.LLM_WITHOUT_ALIGN,
        'hybrid': TranslationMode.HYBRID
    }
    
    print(f"\n{'='*60}")
    print(f"Document Translator")
    print(f"{'='*60}")
    print(f"Input:      {input_path}")
    print(f"Output:     {args.output}")
    print(f"Direction:  {args.source.upper()} → {args.target.upper()}")
    print(f"Mode:       {args.mode}")
    if args.nmt:
        print(f"NMT:        {args.nmt}")
    if args.nmt == 'nllb':
        print(f"NLLB Size:  {args.nllb_size}")
    if args.llm:
        print(f"LLM:        {args.llm}")
    if args.aligner:
        print(f"Aligner:    {args.aligner}")
    print(f"{'='*60}\n")
    
    translator = UltimateDocumentTranslator(
        src_lang=args.source,
        tgt_lang=args.target,
        mode=mode_map[args.mode],
        nmt_backend=args.nmt,
        llm_provider=args.llm,
        aligner=args.aligner,
        nllb_model_size=args.nllb_size
    )
    
    await translator.translate_document(input_path, Path(args.output))
    
    print(f"\n{'='*60}")
    print(f"✓ Success!")
    print(f"Output saved to: {args.output}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    asyncio.run(main())