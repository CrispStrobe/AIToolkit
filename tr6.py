#!/usr/bin/env python3
"""
ULTIMATE DOCUMENT TRANSLATOR
Combines NMT + LLM translation with alignment-based formatting preservation
Supports: Word documents (.docx)
Translation: CTranslate2 (fast) → LLM (fallback)
Alignment: Lindat API → fast_align → SimAlign → heuristic
"""

import argparse
import asyncio
import logging
import os
import random
import re
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Any
import requests
from tqdm import tqdm

# Check and import required libraries
print("🔍 Checking libraries...")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.text.paragraph import Paragraph
    HAS_DOCX = True
    print("  ✓ python-docx")
except ImportError:
    print("  ✗ python-docx (required)")
    sys.exit(1)

# Optional: CTranslate2
try:
    import ctranslate2
    from transformers import AutoTokenizer
    from huggingface_hub import snapshot_download
    HAS_CT2 = True
    print("  ✓ CTranslate2")
except ImportError:
    HAS_CT2 = False
    print("  ⊘ CTranslate2 (optional - will use LLM only)")

# Optional: OpenAI
try:
    from openai import AsyncOpenAI
    HAS_OPENAI = True
    print("  ✓ OpenAI")
except ImportError:
    HAS_OPENAI = False
    print("  ⊘ OpenAI (optional)")

# Optional: fast_align
try:
    from fast_align import align
    HAS_FAST_ALIGN = True
    print("  ✓ fast_align")
except ImportError:
    HAS_FAST_ALIGN = False
    print("  ⊘ fast_align (optional)")

# Optional: simalign
try:
    from simalign import SentenceAligner
    HAS_SIMALIGN = True
    print("  ✓ simalign")
except ImportError:
    HAS_SIMALIGN = False
    print("  ⊘ simalign (optional)")

print("-" * 60)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class FormatRun:
    """A run of text with formatting"""
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    font_color: Optional[str] = None


@dataclass
class TranslatableParagraph:
    """Paragraph with formatting and metadata"""
    runs: List[FormatRun] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def get_text(self) -> str:
        return ''.join(run.text for run in self.runs)
    
    def get_words(self) -> List[str]:
        """Get words for alignment"""
        return self.get_text().split()
    
    def get_formatted_word_indices(self) -> Dict[str, Set[int]]:
        """
        Extract which word indices have which formatting.
        Returns: {'italic': {0, 3}, 'bold': {1}, 'italic_bold': {5}}
        """
        formatted = {'italic': set(), 'bold': set(), 'italic_bold': set()}
        
        text = self.get_text()
        words = text.split()
        
        word_idx = 0
        char_pos = 0
        
        for run in self.runs:
            if not run.text or not run.text.strip():
                continue
            
            # Find words in this run
            run_words = run.text.split()
            
            for word in run_words:
                if word_idx >= len(words):
                    break
                
                # Classify formatting
                if run.bold and run.italic:
                    formatted['italic_bold'].add(word_idx)
                elif run.italic:
                    formatted['italic'].add(word_idx)
                elif run.bold:
                    formatted['bold'].add(word_idx)
                
                word_idx += 1
        
        return formatted


# ============================================================================
# TRANSLATION BACKENDS
# ============================================================================

class CTranslate2Translator:
    """CTranslate2 for fast translation"""
    
    MODELS = {
        'en_to_x': 'cstr/wmt21ct2_int8',
        'x_to_en': 'cstr/wmt21-x-en-ct2-int8',
    }
    
    SUPPORTED_LANGS = ['de', 'es', 'fr', 'it', 'ja', 'zh', 'ru', 'pt', 'nl', 'cs', 'uk']
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = False
        
        if not HAS_CT2:
            return
        
        # Determine direction
        if src_lang == 'en' and tgt_lang in self.SUPPORTED_LANGS:
            self.direction = 'en_to_x'
            self.tokenizer_name = "facebook/wmt21-dense-24-wide-en-x"
        elif tgt_lang == 'en' and src_lang in self.SUPPORTED_LANGS:
            self.direction = 'x_to_en'
            self.tokenizer_name = "facebook/wmt21-dense-24-wide-x-en"
        else:
            logger.info(f"⊘ CT2: {src_lang}→{tgt_lang} not supported")
            return
        
        try:
            logger.info(f"Loading CT2 model...")
            model_path = self._get_or_download_model()
            self.translator = ctranslate2.Translator(model_path, device="auto")
            self.tokenizer = AutoTokenizer.from_pretrained(self.tokenizer_name)
            self.tokenizer.src_lang = src_lang
            self.tokenizer.tgt_lang = tgt_lang
            self.available = True
            logger.info(f"✓ CT2 ready ({src_lang}→{tgt_lang})")
        except Exception as e:
            logger.error(f"CT2 init failed: {e}")
    
    def _get_or_download_model(self) -> str:
        cache_base = Path.home() / '.cache' / 'huggingface' / 'hub'
        model_repo = self.MODELS[self.direction]
        model_dir = cache_base / f"models--{model_repo.replace('/', '--')}"
        ref_path = model_dir / 'refs' / 'main'
        
        if ref_path.exists():
            with open(ref_path) as f:
                commit_hash = f.read().strip()
            snapshot_path = model_dir / 'snapshots' / commit_hash
            if (snapshot_path / 'model.bin').exists():
                return str(snapshot_path)
        
        logger.info("Downloading CT2 model...")
        return snapshot_download(repo_id=model_repo)
    
    def translate_batch(self, texts: List[str]) -> List[str]:
        """Translate batch of texts"""
        if not self.available or not texts:
            return texts
        
        try:
            # Tokenize
            source_batches = []
            for text in texts:
                if not text.strip():
                    source_batches.append([self.tokenizer.src_lang])
                    continue
                
                tokens = self.tokenizer.tokenize(text)
                if self.tokenizer.src_lang not in tokens:
                    tokens = [self.tokenizer.src_lang] + tokens + [self.tokenizer.eos_token]
                source_batches.append(tokens)
            
            # Translate
            target_prefix = [self.tokenizer.lang_code_to_token[self.tgt_lang]]
            results = self.translator.translate_batch(
                source_batches,
                target_prefix=[target_prefix] * len(texts),
                beam_size=5,
                repetition_penalty=1.5
            )
            
            # Decode
            translated = []
            for i, res in enumerate(results):
                if not texts[i].strip():
                    translated.append(texts[i])
                    continue
                
                target_tokens = res.hypotheses[0][len(target_prefix):]
                decoded = self.tokenizer.decode(
                    self.tokenizer.convert_tokens_to_ids(target_tokens),
                    skip_special_tokens=True
                )
                translated.append(decoded.strip())
            
            return translated
        except Exception as e:
            logger.error(f"CT2 translation failed: {e}")
            return texts


class LLMTranslator:
    """LLM translator (GPT-4/Ollama)"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.providers = self._init_providers()
        
        if self.providers:
            logger.info(f"✓ LLM available ({len(self.providers)} providers)")
    
    def _init_providers(self) -> Dict[str, Any]:
        providers = {}
        
        # OpenAI
        if HAS_OPENAI and os.getenv("OPENAI_API_KEY"):
            providers["openai"] = {
                "client": AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY")),
                "model": "gpt-4o-mini"
            }
        
        # Ollama
        if self._check_ollama():
            providers["ollama"] = {
                "url": "http://localhost:11434/api/generate",
                "model": self._get_ollama_model()
            }
        
        return providers
    
    def _check_ollama(self) -> bool:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            return r.status_code == 200 and len(r.json().get("models", [])) > 0
        except:
            return False
    
    def _get_ollama_model(self) -> str:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            models = r.json().get("models", [])
            return models[0]["name"] if models else "llama3.2"
        except:
            return "llama3.2"
    
    async def translate_text(self, text: str) -> Optional[str]:
        """Translate single text"""
        if not text.strip() or not self.providers:
            return None
        
        prompt = f"Translate from {self.src_lang} to {self.tgt_lang}. Return ONLY the translation:\n\n{text}"
        
        # Try OpenAI first
        if "openai" in self.providers:
            try:
                p = self.providers["openai"]
                response = await p["client"].chat.completions.create(
                    model=p["model"],
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                    max_tokens=4000
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                logger.debug(f"OpenAI failed: {e}")
        
        # Try Ollama
        if "ollama" in self.providers:
            try:
                p = self.providers["ollama"]
                r = requests.post(
                    p["url"],
                    json={"model": p["model"], "prompt": prompt, "stream": False},
                    timeout=120
                )
                if r.status_code == 200:
                    return r.json().get("response", "").strip()
            except Exception as e:
                logger.debug(f"Ollama failed: {e}")
        
        return None
    
    async def translate_batch(self, texts: List[str]) -> List[str]:
        """Translate batch"""
        tasks = [self.translate_text(text) for text in texts]
        results = await asyncio.gather(*tasks)
        return [res if res else text for res, text in zip(results, texts)]


# ============================================================================
# WORD ALIGNERS
# ============================================================================

class LindatAligner:
    """Lindat word alignment API"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = self._check_available()
    
    def _check_available(self) -> bool:
        try:
            r = requests.get(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                timeout=5
            )
            return r.status_code in [200, 405]  # 405 = GET not allowed but endpoint exists
        except:
            return False
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            r = requests.post(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                headers={'Content-Type': 'application/json'},
                json={'src_tokens': [src_words], 'trg_tokens': [tgt_words]},
                timeout=30
            )
            
            if r.status_code == 200:
                alignment = r.json()["alignment"][0]
                return [(int(a[0]), int(a[1])) for a in alignment]
        except Exception as e:
            logger.debug(f"Lindat alignment failed: {e}")
        
        return []


class FastAlignAligner:
    """fast_align local aligner"""
    
    def __init__(self):
        self.available = HAS_FAST_ALIGN
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            # fast_align expects "src ||| tgt" format
            src_text = ' '.join(src_words)
            tgt_text = ' '.join(tgt_words)
            alignment_text = f"{src_text} ||| {tgt_text}"
            
            # Run alignment
            result = align([alignment_text], forward=True)
            
            # Parse result
            alignments = []
            for align_str in result[0].split():
                src_idx, tgt_idx = map(int, align_str.split('-'))
                alignments.append((src_idx, tgt_idx))
            
            return alignments
        except Exception as e:
            logger.debug(f"fast_align failed: {e}")
            return []


class SimAlignAligner:
    """SimAlign aligner"""
    
    def __init__(self):
        self.available = HAS_SIMALIGN
        if self.available:
            try:
                self.aligner = SentenceAligner(model="bert", token_type="bpe", matching_methods="mai")
            except Exception as e:
                logger.debug(f"SimAlign init failed: {e}")
                self.available = False
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            result = self.aligner.get_word_aligns(src_words, tgt_words)
            # result['mwmf'] contains alignment dict
            alignments = []
            for src_idx, tgt_idx in result['mwmf'].items():
                alignments.append((src_idx, tgt_idx))
            return alignments
        except Exception as e:
            logger.debug(f"SimAlign failed: {e}")
            return []


class HeuristicAligner:
    """Heuristic fallback - align words that appear in both"""
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Simple heuristic alignment"""
        alignments = []
        
        # Normalize words
        src_lower = [w.lower().strip('.,!?;:') for w in src_words]
        tgt_lower = [w.lower().strip('.,!?;:') for w in tgt_words]
        
        # Find exact matches
        for i, src_word in enumerate(src_lower):
            for j, tgt_word in enumerate(tgt_lower):
                if src_word == tgt_word and len(src_word) > 2:
                    alignments.append((i, j))
        
        return alignments


class MultiAligner:
    """Try multiple aligners in sequence"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.aligners = []
        
        # Initialize all available aligners
        lindat = LindatAligner(src_lang, tgt_lang)
        if lindat.available:
            self.aligners.append(("Lindat", lindat))
            logger.info("✓ Lindat aligner available")
        
        fast_align = FastAlignAligner()
        if fast_align.available:
            self.aligners.append(("fast_align", fast_align))
            logger.info("✓ fast_align available")
        
        simalign = SimAlignAligner()
        if simalign.available:
            self.aligners.append(("SimAlign", simalign))
            logger.info("✓ SimAlign available")
        
        # Always have heuristic fallback
        self.aligners.append(("Heuristic", HeuristicAligner()))
        logger.info("✓ Heuristic aligner (fallback)")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Try aligners in order until one succeeds"""
        for name, aligner in self.aligners:
            result = aligner.align(src_words, tgt_words)
            if result:
                logger.debug(f"Using {name} alignment: {len(result)} links")
                return result
        
        logger.warning("All aligners failed, returning empty alignment")
        return []


# ============================================================================
# DOCUMENT TRANSLATOR
# ============================================================================

class UltimateDocumentTranslator:
    """The ultimate document translator"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        
        # Initialize translators
        self.ct2 = CTranslate2Translator(src_lang, tgt_lang)
        self.llm = LLMTranslator(src_lang, tgt_lang)
        
        # Initialize aligner
        self.aligner = MultiAligner(src_lang, tgt_lang)
        
        # Check we have at least one translator
        if not self.ct2.available and not self.llm.providers:
            logger.error("No translation backends available!")
            sys.exit(1)
    
    async def translate_text(self, text: str) -> str:
        """Translate single text with fallback"""
        if not text.strip():
            return text
        
        # Try CT2 first (fast)
        if self.ct2.available:
            result = self.ct2.translate_batch([text])[0]
            if result and result != text:
                return result
        
        # Fallback to LLM
        if self.llm.providers:
            result = await self.llm.translate_text(text)
            if result:
                return result
        
        logger.error(f"Translation failed for: {text[:50]}...")
        return text
    
    def extract_paragraph(self, para: Paragraph) -> TranslatableParagraph:
        """Extract paragraph with formatting"""
        runs = []
        for run in para.runs:
            runs.append(FormatRun(
                text=run.text,
                bold=run.font.bold,
                italic=run.font.italic,
                underline=run.font.underline,
                font_name=run.font.name,
                font_size=run.font.size.pt if run.font.size else None
            ))
        
        trans_para = TranslatableParagraph(runs=runs)
        trans_para.metadata['style'] = para.style.name if para.style else None
        trans_para.metadata['alignment'] = para.alignment
        
        return trans_para
    
    def apply_aligned_formatting(self, para: Paragraph, trans_para: TranslatableParagraph, 
                                translated_text: str, alignment: List[Tuple[int, int]]):
        """Apply translated text with alignment-based formatting - SAFE VERSION"""
        
        # Get original formatted word indices
        formatted_words = trans_para.get_formatted_word_indices()
        
        # Build target formatting map
        src_words = trans_para.get_words()
        tgt_words = translated_text.split()
        
        if not tgt_words:
            # Empty translation - just clear the paragraph
            para.clear()
            return
        
        tgt_formatting = {}
        
        for src_idx, tgt_idx in alignment:
            if src_idx < len(src_words) and tgt_idx < len(tgt_words):
                for format_type in ['italic_bold', 'bold', 'italic']:
                    if src_idx in formatted_words[format_type]:
                        tgt_formatting[tgt_idx] = format_type
        
        # SAFE APPROACH: Clear runs properly and rebuild
        # First, save paragraph-level properties
        para_style = para.style
        para_alignment = para.alignment
        
        # Clear all runs SAFELY
        while len(para.runs) > 0:
            run_element = para.runs[0]._element
            run_element.getparent().remove(run_element)
        
        # Rebuild with formatting in fewer runs to avoid corruption
        if not alignment or not any(formatted_words.values()):
            # No formatting or alignment - single run
            run = para.add_run(translated_text)
            return
        
        # Group consecutive words with same formatting
        runs_to_create = []
        current_format = tgt_formatting.get(0, None)
        current_words = [tgt_words[0]]
        
        for i in range(1, len(tgt_words)):
            word_format = tgt_formatting.get(i, None)
            
            if word_format == current_format:
                current_words.append(tgt_words[i])
            else:
                # Save current run
                runs_to_create.append((current_format, ' '.join(current_words)))
                # Start new run
                current_format = word_format
                current_words = [tgt_words[i]]
        
        # Add final run
        runs_to_create.append((current_format, ' '.join(current_words)))
        
        # Create runs with spacing
        for i, (format_type, text) in enumerate(runs_to_create):
            # Add space between runs (except first)
            if i > 0:
                para.add_run(' ')
            
            run = para.add_run(text)
            
            # Apply formatting
            if format_type == 'italic_bold':
                run.font.italic = True
                run.font.bold = True
            elif format_type == 'bold':
                run.font.bold = True
            elif format_type == 'italic':
                run.font.italic = True
            
            # Preserve font properties from original
            if trans_para.runs:
                first_orig_run = next((r for r in trans_para.runs if r.text), None)
                if first_orig_run and first_orig_run.font_name:
                    run.font.name = first_orig_run.font_name
        
        # Restore paragraph properties
        try:
            para.style = para_style
        except:
            pass
        para.alignment = para_alignment
    
    def is_paragraph_safe_to_translate(self, para: Paragraph) -> bool:
        """Check if paragraph can be safely translated"""
        # Skip if contains drawings/images
        try:
            for run in para.runs:
                if run._element.xpath('.//w:drawing | .//w:pict'):
                    logger.debug(f"Skipping paragraph with drawing/image")
                    return False
        except:
            pass
        
        # Skip if contains fields (like TOC)
        try:
            if para._element.xpath('.//w:fldChar'):
                logger.debug(f"Skipping paragraph with field")
                return False
        except:
            pass
        
        return True
    
    async def translate_paragraph(self, para: Paragraph):
        """Translate single paragraph - SAFE VERSION"""
        if not para.text or not para.text.strip():
            return
        
        # Safety check
        if not self.is_paragraph_safe_to_translate(para):
            return
        
        try:
            # Extract
            trans_para = self.extract_paragraph(para)
            original_text = trans_para.get_text()
            src_words = trans_para.get_words()
            
            if not src_words:
                return
            
            # Translate
            translated_text = await self.translate_text(original_text)
            tgt_words = translated_text.split()
            
            if not tgt_words:
                return
            
            # Align
            alignment = self.aligner.align(src_words, tgt_words)
            
            # Apply
            self.apply_aligned_formatting(para, trans_para, translated_text, alignment)
            
        except Exception as e:
            logger.error(f"Failed to translate paragraph: {e}")
            # Leave paragraph unchanged on error
    
    def get_all_paragraphs(self, doc: Document) -> List[Paragraph]:
        """Get all paragraphs including tables, headers, footers, footnotes"""
        all_paras = list(doc.paragraphs)
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)
        
        # Headers/footers
        for section in doc.sections:
            all_paras.extend(section.header.paragraphs)
            all_paras.extend(section.footer.paragraphs)
        
        # Footnotes
        try:
            if hasattr(doc, 'part') and hasattr(doc.part, 'footnotes_part'):
                footnotes = doc.part.footnotes_part.footnotes
                for footnote in footnotes:
                    all_paras.extend(footnote.paragraphs)
        except Exception as e:
            logger.debug(f"Could not extract footnotes: {e}")
        
        return all_paras
    
    async def translate_document(self, input_path: Path, output_path: Path):
        """Translate entire document - SAFE VERSION"""
        logger.info(f"Loading document: {input_path}")
        
        try:
            doc = Document(str(input_path))
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            return
        
        # Get all paragraphs
        all_paras = self.get_all_paragraphs(doc)
        translatable = [p for p in all_paras if p.text and p.text.strip() and self.is_paragraph_safe_to_translate(p)]
        
        logger.info(f"Found {len(translatable)} paragraphs to translate")
        
        # Translate with progress bar
        for para in tqdm(translatable, desc="Translating"):
            try:
                await self.translate_paragraph(para)
                await asyncio.sleep(0.1)  # Rate limiting
            except Exception as e:
                logger.error(f"Error translating paragraph: {e}")
                continue
        
        # Save with error handling
        try:
            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Remove output file if it exists
            if output_path.exists():
                output_path.unlink()
            
            # Save
            logger.info(f"Saving to: {output_path}")
            doc.save(str(output_path))
            
            # Verify file was created and is readable
            if not output_path.exists():
                raise Exception("Output file was not created")
            
            if output_path.stat().st_size == 0:
                raise Exception("Output file is empty")
            
            # Try to re-open to verify integrity
            try:
                test_doc = Document(str(output_path))
                logger.info(f"✓ Document verified ({len(test_doc.paragraphs)} paragraphs)")
            except Exception as e:
                raise Exception(f"Output document is corrupted: {e}")
            
            logger.info("✓ Translation complete!")
            
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            if output_path.exists():
                output_path.unlink()
            raise


# ============================================================================
# CLI
# ============================================================================

async def main():
    parser = argparse.ArgumentParser(
        description='Ultimate Document Translator - NMT + LLM + Multi-Aligner'
    )
    
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('output', help='Output .docx file')
    parser.add_argument('-s', '--source', default='en', help='Source language (default: en)')
    parser.add_argument('-t', '--target', default='de', help='Target language (default: de)')
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"Error: Only .docx files supported")
        sys.exit(1)
    
    print(f"\n{'='*60}")
    print(f"Ultimate Document Translator")
    print(f"{'='*60}")
    print(f"Input:  {input_path}")
    print(f"Output: {args.output}")
    print(f"Direction: {args.source.upper()} → {args.target.upper()}")
    print(f"{'='*60}\n")
    
    translator = UltimateDocumentTranslator(args.source, args.target)
    await translator.translate_document(input_path, Path(args.output))
    
    print(f"\n{'='*60}")
    print(f"✓ Success!")
    print(f"Output saved to: {args.output}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    asyncio.run(main())