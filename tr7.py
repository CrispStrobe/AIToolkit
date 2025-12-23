#!/usr/bin/env python3
import argparse
import asyncio
import logging
import os
import re
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Any
import requests
from tqdm import tqdm

# Formatting Libraries
try:
    from docx import Document
    from docx.shared import Pt
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

# Logic for punctuation-aware splitting
TOKEN_RE = re.compile(r'\w+|[^\w\s]')

@dataclass
class FormatRun:
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None

@dataclass
class TranslatableParagraph:
    runs: List[FormatRun] = field(default_factory=list)
    
    def get_text(self) -> str:
        return ''.join(run.text for run in self.runs)
    
    def get_tokens(self) -> List[str]:
        return TOKEN_RE.findall(self.get_text())

    def get_formatted_token_indices(self) -> Dict[str, Set[int]]:
        formatted = {'italic': set(), 'bold': set(), 'italic_bold': set()}
        tokens = self.get_tokens()
        current_char_pos = 0
        
        # Map tokens to their formatting based on character overlap in runs
        for token_idx, token in enumerate(tokens):
            # Find which run contains the start of this token
            start_in_text = self.get_text().find(token, current_char_pos)
            if start_in_text == -1: continue
            
            # Check formatting of the run at this position
            run_char_count = 0
            for run in self.runs:
                run_char_count += len(run.text)
                if run_char_count > start_in_text:
                    if run.bold and run.italic: formatted['italic_bold'].add(token_idx)
                    elif run.bold: formatted['bold'].add(token_idx)
                    elif run.italic: formatted['italic'].add(token_idx)
                    break
            current_char_pos = start_in_text + len(token)
            
        return formatted

# ============================================================================
# TRANSLATION & ALIGNMENT (Simplified for brevity, use your existing backends)
# ============================================================================

# [Your CTranslate2Translator and LLMTranslator classes go here...]
# [Ensure they are integrated into the MultiAligner class as before]

# ============================================================================
# IMPROVED DOCUMENT TRANSLATOR
# ============================================================================

class UltimateDocumentTranslator:
    def __init__(self, src_lang, tgt_lang):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        # Use your existing init logic here...
        from __main__ import CTranslate2Translator, LLMTranslator, MultiAligner
        self.ct2 = CTranslate2Translator(src_lang, tgt_lang)
        self.llm = LLMTranslator(src_lang, tgt_lang)
        self.aligner = MultiAligner(src_lang, tgt_lang)

    async def translate_text(self, text: str) -> str:
        if not text.strip(): return text
        if self.ct2.available:
            res = self.ct2.translate_batch([text])[0]
            if res: return res
        return await self.llm.translate_text(text) if self.llm.providers else text

    def apply_aligned_formatting(self, para: Paragraph, trans_para: TranslatableParagraph, 
                                translated_text: str, alignment: List[Tuple[int, int]]):
        """Fixes 'Word Trash' by handling punctuation and spacing properly"""
        formatted_indices = trans_para.get_formatted_token_indices()
        src_tokens = trans_para.get_tokens()
        tgt_tokens = TOKEN_RE.findall(translated_text)
        
        if not tgt_tokens: return

        # Map formatting to target tokens
        tgt_formatting = {}
        for src_idx, tgt_idx in alignment:
            if src_idx < len(src_tokens) and tgt_idx < len(tgt_tokens):
                for ftype in ['italic_bold', 'bold', 'italic']:
                    if src_idx in formatted_indices[ftype]:
                        tgt_formatting[tgt_idx] = ftype

        # Clear existing runs
        p_element = para._element
        for run in para.runs:
            p_element.remove(run._element)

        # Rebuild runs with smart spacing
        for i, token in enumerate(tgt_tokens):
            run = para.add_run(token)
            
            # Apply Style
            ftype = tgt_formatting.get(i)
            if ftype == 'italic_bold': run.italic = run.bold = True
            elif ftype == 'bold': run.bold = True
            elif ftype == 'italic': run.italic = True

            # Smart Spacing: Add space if the next token is not punctuation
            if i < len(tgt_tokens) - 1:
                next_token = tgt_tokens[i+1]
                if next_token not in ".,!?;:)]}»" and token not in "([{«":
                    para.add_run(" ")

    def is_paragraph_safe(self, para: Paragraph) -> bool:
        if not para.text.strip(): return False
        # Allow paragraphs with footnote references
        xml = para._element.xml
        if 'w:drawing' in xml or 'w:pict' in xml: return False
        return True

    async def process_document(self, input_path: Path, output_path: Path):
        doc = Document(str(input_path))
        
        # Gather all sections including Footnotes
        targets = []
        for p in doc.paragraphs: targets.append(p)
        
        # Deep XML Footnote extraction
        try:
            footnotes_part = doc.part.footnotes_part
            if footnotes_part:
                for footnote in footnotes_part.footnotes:
                    for p in footnote.paragraphs:
                        targets.append(p)
        except: pass

        for para in tqdm(targets, desc="Translating"):
            if not self.is_paragraph_safe(para): continue
            
            # Translation logic
            trans_para = TranslatableParagraph([FormatRun(r.text, r.bold, r.italic) for r in para.runs])
            translated = await self.translate_text(trans_para.get_text())
            
            # Align tokens
            alignment = self.aligner.align(trans_para.get_tokens(), TOKEN_RE.findall(translated))
            
            # Apply
            self.apply_aligned_formatting(para, trans_para, translated, alignment)

        doc.save(str(output_path))