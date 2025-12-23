#!/usr/bin/env python3
"""
ULTIMATE DOCUMENT TRANSLATOR - FIXED VERSION
Fixes: footnote translation, formatting alignment, spacing issues
"""

import argparse
import asyncio
import logging
import os
import random
import re
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Any
import requests
from tqdm import tqdm

# Check and import required libraries
print("🔍 Checking libraries...")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.text.paragraph import Paragraph
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    HAS_DOCX = True
    print("  ✓ python-docx")
except ImportError:
    print("  ✗ python-docx (required)")
    sys.exit(1)

# Optional: CTranslate2
try:
    import ctranslate2
    from transformers import AutoTokenizer
    from huggingface_hub import snapshot_download
    HAS_CT2 = True
    print("  ✓ CTranslate2")
except ImportError:
    HAS_CT2 = False
    print("  ⊘ CTranslate2 (optional - will use LLM only)")

# Optional: OpenAI
try:
    from openai import AsyncOpenAI
    HAS_OPENAI = True
    print("  ✓ OpenAI")
except ImportError:
    HAS_OPENAI = False
    print("  ⊘ OpenAI (optional)")

# Optional: fast_align
try:
    from fast_align import align
    HAS_FAST_ALIGN = True
    print("  ✓ fast_align")
except ImportError:
    HAS_FAST_ALIGN = False
    print("  ⊘ fast_align (optional)")

# Optional: simalign
try:
    from simalign import SentenceAligner
    HAS_SIMALIGN = True
    print("  ✓ simalign")
except ImportError:
    HAS_SIMALIGN = False
    print("  ⊘ simalign (optional)")

print("-" * 60)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class FormatRun:
    """A run of text with formatting"""
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    font_color: Optional[str] = None


@dataclass
class TranslatableParagraph:
    """Paragraph with formatting and metadata"""
    runs: List[FormatRun] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def get_text(self) -> str:
        return ''.join(run.text for run in self.runs)
    
    def get_words(self) -> List[str]:
        """Get words for alignment"""
        return self.get_text().split()


# ============================================================================
# TRANSLATION BACKENDS
# ============================================================================

class CTranslate2Translator:
    """CTranslate2 for fast translation"""
    
    MODELS = {
        'en_to_x': 'cstr/wmt21ct2_int8',
        'x_to_en': 'cstr/wmt21-x-en-ct2-int8',
    }
    
    SUPPORTED_LANGS = ['de', 'es', 'fr', 'it', 'ja', 'zh', 'ru', 'pt', 'nl', 'cs', 'uk']
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = False
        
        if not HAS_CT2:
            return
        
        # Determine direction
        if src_lang == 'en' and tgt_lang in self.SUPPORTED_LANGS:
            self.direction = 'en_to_x'
            self.tokenizer_name = "facebook/wmt21-dense-24-wide-en-x"
        elif tgt_lang == 'en' and src_lang in self.SUPPORTED_LANGS:
            self.direction = 'x_to_en'
            self.tokenizer_name = "facebook/wmt21-dense-24-wide-x-en"
        else:
            logger.info(f"⊘ CT2: {src_lang}→{tgt_lang} not supported")
            return
        
        try:
            logger.info(f"Loading CT2 model...")
            model_path = self._get_or_download_model()
            self.translator = ctranslate2.Translator(model_path, device="auto")
            self.tokenizer = AutoTokenizer.from_pretrained(self.tokenizer_name)
            self.tokenizer.src_lang = src_lang
            self.tokenizer.tgt_lang = tgt_lang
            self.available = True
            logger.info(f"✓ CT2 ready ({src_lang}→{tgt_lang})")
        except Exception as e:
            logger.error(f"CT2 init failed: {e}")
    
    def _get_or_download_model(self) -> str:
        cache_base = Path.home() / '.cache' / 'huggingface' / 'hub'
        model_repo = self.MODELS[self.direction]
        model_dir = cache_base / f"models--{model_repo.replace('/', '--')}"
        ref_path = model_dir / 'refs' / 'main'
        
        if ref_path.exists():
            with open(ref_path) as f:
                commit_hash = f.read().strip()
            snapshot_path = model_dir / 'snapshots' / commit_hash
            if (snapshot_path / 'model.bin').exists():
                return str(snapshot_path)
        
        logger.info("Downloading CT2 model...")
        return snapshot_download(repo_id=model_repo)
    
    def translate_batch(self, texts: List[str]) -> List[str]:
        """Translate batch of texts"""
        if not self.available or not texts:
            return texts
        
        try:
            # Tokenize
            source_batches = []
            for text in texts:
                if not text.strip():
                    source_batches.append([self.tokenizer.src_lang])
                    continue
                
                tokens = self.tokenizer.tokenize(text)
                if self.tokenizer.src_lang not in tokens:
                    tokens = [self.tokenizer.src_lang] + tokens + [self.tokenizer.eos_token]
                source_batches.append(tokens)
            
            # Translate
            target_prefix = [self.tokenizer.lang_code_to_token[self.tgt_lang]]
            results = self.translator.translate_batch(
                source_batches,
                target_prefix=[target_prefix] * len(texts),
                beam_size=5,
                repetition_penalty=1.5
            )
            
            # Decode
            translated = []
            for i, res in enumerate(results):
                if not texts[i].strip():
                    translated.append(texts[i])
                    continue
                
                target_tokens = res.hypotheses[0][len(target_prefix):]
                decoded = self.tokenizer.decode(
                    self.tokenizer.convert_tokens_to_ids(target_tokens),
                    skip_special_tokens=True
                )
                translated.append(decoded.strip())
            
            return translated
        except Exception as e:
            logger.error(f"CT2 translation failed: {e}")
            return texts


class LLMTranslator:
    """LLM translator (GPT-4/Ollama)"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.providers = self._init_providers()
        
        if self.providers:
            logger.info(f"✓ LLM available ({len(self.providers)} providers)")
    
    def _init_providers(self) -> Dict[str, Any]:
        providers = {}
        
        # OpenAI
        if HAS_OPENAI and os.getenv("OPENAI_API_KEY"):
            providers["openai"] = {
                "client": AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY")),
                "model": "gpt-4o-mini"
            }
        
        # Ollama
        if self._check_ollama():
            providers["ollama"] = {
                "url": "http://localhost:11434/api/generate",
                "model": self._get_ollama_model()
            }
        
        return providers
    
    def _check_ollama(self) -> bool:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            return r.status_code == 200 and len(r.json().get("models", [])) > 0
        except:
            return False
    
    def _get_ollama_model(self) -> str:
        try:
            r = requests.get("http://localhost:11434/api/tags", timeout=2)
            models = r.json().get("models", [])
            return models[0]["name"] if models else "llama3.2"
        except:
            return "llama3.2"
    
    async def translate_text(self, text: str) -> Optional[str]:
        """Translate single text"""
        if not text.strip() or not self.providers:
            return None
        
        prompt = f"Translate from {self.src_lang} to {self.tgt_lang}. Return ONLY the translation:\n\n{text}"
        
        # Try OpenAI first
        if "openai" in self.providers:
            try:
                p = self.providers["openai"]
                response = await p["client"].chat.completions.create(
                    model=p["model"],
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                    max_tokens=4000
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                logger.debug(f"OpenAI failed: {e}")
        
        # Try Ollama
        if "ollama" in self.providers:
            try:
                p = self.providers["ollama"]
                r = requests.post(
                    p["url"],
                    json={"model": p["model"], "prompt": prompt, "stream": False},
                    timeout=120
                )
                if r.status_code == 200:
                    return r.json().get("response", "").strip()
            except Exception as e:
                logger.debug(f"Ollama failed: {e}")
        
        return None
    
    async def translate_batch(self, texts: List[str]) -> List[str]:
        """Translate batch"""
        tasks = [self.translate_text(text) for text in texts]
        results = await asyncio.gather(*tasks)
        return [res if res else text for res, text in zip(results, texts)]


# ============================================================================
# WORD ALIGNERS
# ============================================================================

class LindatAligner:
    """Lindat word alignment API"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.available = self._check_available()
    
    def _check_available(self) -> bool:
        try:
            r = requests.get(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                timeout=5
            )
            return r.status_code in [200, 405]
        except:
            return False
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            r = requests.post(
                f"https://lindat.cz/services/text-aligner/align/{self.src_lang}-{self.tgt_lang}",
                headers={'Content-Type': 'application/json'},
                json={'src_tokens': [src_words], 'trg_tokens': [tgt_words]},
                timeout=30
            )
            
            if r.status_code == 200:
                alignment = r.json()["alignment"][0]
                return [(int(a[0]), int(a[1])) for a in alignment]
        except Exception as e:
            logger.debug(f"Lindat alignment failed: {e}")
        
        return []


class FastAlignAligner:
    """fast_align local aligner"""
    
    def __init__(self):
        self.available = HAS_FAST_ALIGN
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            src_text = ' '.join(src_words)
            tgt_text = ' '.join(tgt_words)
            alignment_text = f"{src_text} ||| {tgt_text}"
            
            result = align([alignment_text], forward=True)
            
            alignments = []
            for align_str in result[0].split():
                src_idx, tgt_idx = map(int, align_str.split('-'))
                alignments.append((src_idx, tgt_idx))
            
            return alignments
        except Exception as e:
            logger.debug(f"fast_align failed: {e}")
            return []


class SimAlignAligner:
    """SimAlign aligner"""
    
    def __init__(self):
        self.available = HAS_SIMALIGN
        if self.available:
            try:
                self.aligner = SentenceAligner(model="bert", token_type="bpe", matching_methods="mai")
            except Exception as e:
                logger.debug(f"SimAlign init failed: {e}")
                self.available = False
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Align word indices"""
        if not self.available or not src_words or not tgt_words:
            return []
        
        try:
            result = self.aligner.get_word_aligns(src_words, tgt_words)
            alignments = []
            for src_idx, tgt_idx in result['mwmf'].items():
                alignments.append((src_idx, tgt_idx))
            return alignments
        except Exception as e:
            logger.debug(f"SimAlign failed: {e}")
            return []


class HeuristicAligner:
    """Heuristic fallback - align words that appear in both"""
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Simple heuristic alignment"""
        alignments = []
        
        src_lower = [w.lower().strip('.,!?;:') for w in src_words]
        tgt_lower = [w.lower().strip('.,!?;:') for w in tgt_words]
        
        for i, src_word in enumerate(src_lower):
            for j, tgt_word in enumerate(tgt_lower):
                if src_word == tgt_word and len(src_word) > 2:
                    alignments.append((i, j))
        
        return alignments


class MultiAligner:
    """Try multiple aligners in sequence"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.aligners = []
        
        lindat = LindatAligner(src_lang, tgt_lang)
        if lindat.available:
            self.aligners.append(("Lindat", lindat))
            logger.info("✓ Lindat aligner available")
        
        fast_align = FastAlignAligner()
        if fast_align.available:
            self.aligners.append(("fast_align", fast_align))
            logger.info("✓ fast_align available")
        
        simalign = SimAlignAligner()
        if simalign.available:
            self.aligners.append(("SimAlign", simalign))
            logger.info("✓ SimAlign available")
        
        self.aligners.append(("Heuristic", HeuristicAligner()))
        logger.info("✓ Heuristic aligner (fallback)")
    
    def align(self, src_words: List[str], tgt_words: List[str]) -> List[Tuple[int, int]]:
        """Try aligners in order until one succeeds"""
        for name, aligner in self.aligners:
            result = aligner.align(src_words, tgt_words)
            if result:
                logger.debug(f"Using {name} alignment: {len(result)} links")
                return result
        
        logger.warning("All aligners failed, returning empty alignment")
        return []


# ============================================================================
# DOCUMENT TRANSLATOR - FIXED VERSION
# ============================================================================

class UltimateDocumentTranslator:
    """The ultimate document translator - FIXED"""
    
    def __init__(self, src_lang: str, tgt_lang: str):
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        
        # Initialize translators
        self.ct2 = CTranslate2Translator(src_lang, tgt_lang)
        self.llm = LLMTranslator(src_lang, tgt_lang)
        
        # Initialize aligner
        self.aligner = MultiAligner(src_lang, tgt_lang)
        
        if not self.ct2.available and not self.llm.providers:
            logger.error("No translation backends available!")
            sys.exit(1)
    
    async def translate_text(self, text: str) -> str:
        """Translate single text with fallback"""
        if not text.strip():
            return text
        
        # Try CT2 first (fast)
        if self.ct2.available:
            result = self.ct2.translate_batch([text])[0]
            if result and result != text:
                return result
        
        # Fallback to LLM
        if self.llm.providers:
            result = await self.llm.translate_text(text)
            if result:
                return result
        
        logger.error(f"Translation failed for: {text[:50]}...")
        return text
    
    def extract_paragraph(self, para: Paragraph) -> TranslatableParagraph:
        """Extract paragraph with formatting"""
        runs = []
        for run in para.runs:
            runs.append(FormatRun(
                text=run.text,
                bold=run.font.bold,
                italic=run.font.italic,
                underline=run.font.underline,
                font_name=run.font.name,
                font_size=run.font.size.pt if run.font.size else None
            ))
        
        trans_para = TranslatableParagraph(runs=runs)
        trans_para.metadata['style'] = para.style.name if para.style else None
        trans_para.metadata['alignment'] = para.alignment
        
        return trans_para
    
    def apply_formatting_runbased(self, para: Paragraph, trans_para: TranslatableParagraph, 
                                   translated_text: str):
        """
        SIMPLIFIED: Apply formatting by preserving run structure.
        Instead of word alignment, translate each formatted segment separately.
        """
        # Save paragraph-level properties
        para_style = para.style
        para_alignment = para.alignment
        
        # Clear all runs
        while len(para.runs) > 0:
            run_element = para.runs[0]._element
            run_element.getparent().remove(run_element)
        
        # Simple case: just add translated text
        run = para.add_run(translated_text)
        
        # Try to preserve font from first non-empty run
        for orig_run in trans_para.runs:
            if orig_run.text and orig_run.text.strip():
                if orig_run.font_name:
                    run.font.name = orig_run.font_name
                if orig_run.font_size:
                    run.font.size = Pt(orig_run.font_size)
                break
        
        # Restore paragraph properties
        try:
            para.style = para_style
        except:
            pass
        para.alignment = para_alignment
    
    def get_formatted_segments(self, trans_para: TranslatableParagraph) -> List[Tuple[str, bool, bool]]:
        """
        Extract text segments with their formatting.
        Returns: [(text, is_bold, is_italic), ...]
        """
        segments = []
        for run in trans_para.runs:
            if run.text:
                segments.append((
                    run.text,
                    run.bold if run.bold is not None else False,
                    run.italic if run.italic is not None else False
                ))
        return segments
    
    async def translate_with_formatting(self, para: Paragraph, trans_para: TranslatableParagraph):
        """
        IMPROVED: Translate segments individually to preserve formatting better.
        """
        segments = self.get_formatted_segments(trans_para)
        
        if not segments:
            return
        
        # If no formatting differences, translate as one piece
        if all(bold == segments[0][1] and italic == segments[0][2] 
               for _, bold, italic in segments):
            translated = await self.translate_text(trans_para.get_text())
            self.apply_simple_formatting(para, trans_para, translated,
                                         segments[0][1], segments[0][2])
            return
        
        # Save paragraph properties
        para_style = para.style
        para_alignment = para.alignment
        
        # Clear runs
        while len(para.runs) > 0:
            run_element = para.runs[0]._element
            run_element.getparent().remove(run_element)
        
        # Translate and apply each segment
        for i, (text, is_bold, is_italic) in enumerate(segments):
            if not text.strip():
                # Preserve whitespace runs as-is
                run = para.add_run(text)
                continue
            
            # Translate this segment
            translated_segment = await self.translate_text(text)
            
            # Create run with formatting
            run = para.add_run(translated_segment)
            run.font.bold = is_bold
            run.font.italic = is_italic
            
            # Preserve font properties
            for orig_run in trans_para.runs:
                if orig_run.text and orig_run.text.strip():
                    if orig_run.font_name:
                        run.font.name = orig_run.font_name
                    if orig_run.font_size:
                        run.font.size = Pt(orig_run.font_size)
                    break
        
        # Restore paragraph properties
        try:
            para.style = para_style
        except:
            pass
        para.alignment = para_alignment
    
    def apply_simple_formatting(self, para: Paragraph, trans_para: TranslatableParagraph,
                               translated_text: str, is_bold: bool, is_italic: bool):
        """Apply simple formatting to entire paragraph"""
        para_style = para.style
        para_alignment = para.alignment
        
        # Clear runs
        while len(para.runs) > 0:
            run_element = para.runs[0]._element
            run_element.getparent().remove(run_element)
        
        # Add translated text
        run = para.add_run(translated_text)
        run.font.bold = is_bold
        run.font.italic = is_italic
        
        # Preserve font
        for orig_run in trans_para.runs:
            if orig_run.text and orig_run.text.strip():
                if orig_run.font_name:
                    run.font.name = orig_run.font_name
                if orig_run.font_size:
                    run.font.size = Pt(orig_run.font_size)
                break
        
        # Restore paragraph properties
        try:
            para.style = para_style
        except:
            pass
        para.alignment = para_alignment
    
    def is_paragraph_safe_to_translate(self, para: Paragraph) -> bool:
        """Check if paragraph can be safely translated"""
        if not para.text or not para.text.strip():
            return False
        
        # Skip if contains drawings/images
        try:
            for run in para.runs:
                if run._element.xpath('.//w:drawing | .//w:pict'):
                    logger.debug(f"Skipping paragraph with drawing/image")
                    return False
        except:
            pass
        
        # Skip if contains non-footnote fields
        try:
            field_chars = para._element.xpath('.//w:fldChar')
            if field_chars:
                is_footnote_field = para._element.xpath('.//w:footnoteReference')
                if not is_footnote_field:
                    logger.debug(f"Skipping paragraph with field")
                    return False
        except:
            pass
        
        return True
    
    async def translate_paragraph(self, para: Paragraph):
        """Translate single paragraph - IMPROVED VERSION"""
        if not para.text or not para.text.strip():
            return
        
        if not self.is_paragraph_safe_to_translate(para):
            return
        
        try:
            trans_para = self.extract_paragraph(para)
            await self.translate_with_formatting(para, trans_para)
        except Exception as e:
            logger.error(f"Failed to translate paragraph: {e}")
            logger.debug(f"  Text: {para.text[:100]}...")
    
    def get_footnotes(self, doc: Document) -> List[Paragraph]:
        """
        FIXED: Better footnote extraction using XML parsing
        """
        footnotes = []
        
        try:
            # Access document's XML parts
            document_part = doc.part
            
            # Get footnotes part
            try:
                footnotes_part_rId = None
                for rel in document_part.rels.values():
                    if "footnotes" in rel.target_ref:
                        footnotes_part = rel.target_part
                        break
                else:
                    return footnotes
                
                # Parse footnotes XML
                footnotes_xml = footnotes_part._blob
                from lxml import etree
                root = etree.fromstring(footnotes_xml)
                
                # Find all footnote elements
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                for footnote_elem in root.findall('.//w:footnote', namespaces):
                    footnote_id = footnote_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    
                    # Skip special footnotes (separator, continuation)
                    if footnote_id in ['-1', '0']:
                        continue
                    
                    # Extract paragraphs from footnote
                    for para_elem in footnote_elem.findall('.//w:p', namespaces):
                        # Create a temporary paragraph object
                        para = Paragraph(para_elem, document_part)
                        if para.text and para.text.strip():
                            footnotes.append(para)
                
                logger.info(f"Extracted {len(footnotes)} footnote paragraphs")
                
            except Exception as e:
                logger.warning(f"Could not extract footnotes: {e}")
        
        except Exception as e:
            logger.warning(f"Footnote extraction failed: {e}")
        
        return footnotes
    
    def get_all_paragraphs(self, doc: Document) -> List[Tuple[Paragraph, str]]:
        """
        Get all paragraphs with location tags.
        Returns: List of (paragraph, location_type) tuples
        """
        all_paras = []
        
        # Main document
        for para in doc.paragraphs:
            all_paras.append((para, "body"))
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_paras.append((para, "table"))
        
        # Headers/footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                all_paras.append((para, "header"))
            for para in section.footer.paragraphs:
                all_paras.append((para, "footer"))
        
        # Footnotes - using improved extraction
        footnote_paras = self.get_footnotes(doc)
        for para in footnote_paras:
            all_paras.append((para, "footnote"))
        
        return all_paras
    
    async def translate_document(self, input_path: Path, output_path: Path):
        """Translate entire document - FIXED VERSION"""
        logger.info(f"Loading document: {input_path}")
        
        try:
            doc = Document(str(input_path))
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            return
        
        # Get all paragraphs
        all_paras_with_location = self.get_all_paragraphs(doc)
        
        # Filter translatable
        translatable = []
        for para, location in all_paras_with_location:
            if para.text and para.text.strip() and self.is_paragraph_safe_to_translate(para):
                translatable.append((para, location))
        
        # Log statistics
        by_location = defaultdict(int)
        for _, location in translatable:
            by_location[location] += 1
        
        logger.info(f"Found {len(translatable)} paragraphs to translate:")
        for location, count in by_location.items():
            logger.info(f"  - {location}: {count}")
        
        # Translate with progress bar
        for para, location in tqdm(translatable, desc="Translating"):
            try:
                await self.translate_paragraph(para)
                await asyncio.sleep(0.1)
            except Exception as e:
                logger.error(f"Error translating {location} paragraph: {e}")
                logger.debug(f"  Text: {para.text[:100]}...")
                continue
        
        # Save
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            if output_path.exists():
                output_path.unlink()
            
            logger.info(f"Saving to: {output_path}")
            doc.save(str(output_path))
            
            if not output_path.exists():
                raise Exception("Output file was not created")
            
            if output_path.stat().st_size == 0:
                raise Exception("Output file is empty")
            
            # Verify
            try:
                test_doc = Document(str(output_path))
                logger.info(f"✓ Document verified ({len(test_doc.paragraphs)} paragraphs)")
            except Exception as e:
                raise Exception(f"Output document is corrupted: {e}")
            
            logger.info("✓ Translation complete!")
            
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            if output_path.exists():
                output_path.unlink()
            raise


# ============================================================================
# CLI
# ============================================================================

async def main():
    parser = argparse.ArgumentParser(
        description='Ultimate Document Translator - FIXED - NMT + LLM'
    )
    
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('output', help='Output .docx file')
    parser.add_argument('-s', '--source', default='en', help='Source language (default: en)')
    parser.add_argument('-t', '--target', default='de', help='Target language (default: de)')
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"Error: Only .docx files supported")
        sys.exit(1)
    
    print(f"\n{'='*60}")
    print(f"Ultimate Document Translator")
    print(f"{'='*60}")
    print(f"Input:  {input_path}")
    print(f"Output: {args.output}")
    print(f"Direction: {args.source.upper()} → {args.target.upper()}")
    print(f"{'='*60}\n")
    
    translator = UltimateDocumentTranslator(args.source, args.target)
    await translator.translate_document(input_path, Path(args.output))
    
    print(f"\n{'='*60}")
    print(f"✓ Success!")
    print(f"Output saved to: {args.output}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    asyncio.run(main())